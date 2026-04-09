# app.py — Flask server untuk News Scraper
import os, json, csv, io, threading, re, unicodedata, logging, time
import numpy as np
from functools import wraps
import pdfplumber
import fitz  # PyMuPDF
from datetime import datetime, date, timedelta
from flask import Flask, render_template, request, jsonify, Response, g
import hashlib
from flask_cors import CORS
from apscheduler.schedulers.background import BackgroundScheduler
from apscheduler.triggers.cron import CronTrigger

import openai
from scraper import scrape_all, auto_detect_selectors
from kemlu_scraper import is_kemlu_url
from ai_services import generate_ai_summary, check_openai_available, ocr_arabic_page, ocr_arabic_pages_batch
from db_services import push_kb_articles, fetch_kb_articles_from_db, check_supabase_available
from kb_processor import generate_slug, generate_summary, generate_tags, convert_to_kb_format

# ─── Logging Setup ─────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 500 * 1024 * 1024  # 500 MB — cukup untuk PDF 500+ halaman
CORS(app, origins="*")

# ─── Auth ───────────────────────────────────────────────────────────────────────
_ADMIN_USERNAME = os.environ.get("ADMIN_USERNAME", "admin")
_ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "")
_SESSION_SECRET = os.environ.get("SESSION_SECRET", "")

def _hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

def _make_user_token(username: str) -> str:
    return hashlib.sha256(f"{_SESSION_SECRET}:{username}".encode()).hexdigest()

def _load_users() -> list:
    """
    Baca users dari Supabase (primary) → fallback JSON lokal.
    Jika Supabase berhasil, sinkronkan ke JSON lokal juga.
    """
    try:
        from db_services import fetch_app_users
        sb_users = fetch_app_users()
        if sb_users:
            # Sinkronkan ke JSON lokal agar konsisten
            try:
                os.makedirs(DATA_DIR, exist_ok=True)
                with open(USERS_FILE, "w") as f:
                    json.dump(sb_users, f, indent=2)
            except Exception:
                pass
            return sb_users
    except Exception:
        pass
    # Fallback ke JSON lokal
    try:
        if os.path.exists(USERS_FILE):
            with open(USERS_FILE) as f:
                return json.load(f)
    except Exception:
        pass
    return []

def _save_users(users: list) -> None:
    """Simpan ke JSON lokal. Operasi individual (add/delete/reset) juga memanggil Supabase langsung."""
    os.makedirs(DATA_DIR, exist_ok=True)
    with open(USERS_FILE, "w") as f:
        json.dump(users, f, indent=2)

def _get_valid_tokens() -> dict:
    tokens = {_SESSION_SECRET: (_ADMIN_USERNAME, True)}
    for user in _load_users():
        tok = _make_user_token(user["username"])
        tokens[tok] = (user["username"], False)
    return tokens

_PUBLIC_ENDPOINTS = {"login", "static"}

@app.before_request
def _check_auth():
    if request.method == "OPTIONS":
        return
    if request.endpoint in _PUBLIC_ENDPOINTS:
        return
    if not _SESSION_SECRET:
        return jsonify({"error": "Auth not configured. Set SESSION_SECRET env var."}), 500
    auth = request.headers.get("Authorization", "")
    token = auth[7:] if auth.startswith("Bearer ") else ""
    valid = _get_valid_tokens()
    if token not in valid:
        return jsonify({"error": "Unauthorized"}), 401
    g.current_user, g.is_admin = valid[token]
    _update_last_seen(g.current_user)

@app.route("/api/login", methods=["POST"])
def login():
    data = request.get_json(silent=True) or {}
    username = data.get("username", "").strip()
    password = data.get("password", "").strip()
    if not _ADMIN_PASSWORD:
        return jsonify({"error": "Auth not configured. Set ADMIN_PASSWORD secret."}), 500
    if username == _ADMIN_USERNAME and password == _ADMIN_PASSWORD:
        _record_login(username)
        return jsonify({"token": _SESSION_SECRET, "is_admin": True, "username": username})
    users = _load_users()
    pw_hash = _hash_password(password)
    for user in users:
        if user["username"] == username and user["password_hash"] == pw_hash:
            _record_login(username)
            return jsonify({"token": _make_user_token(username), "is_admin": False, "username": username})
    return jsonify({"error": "Username atau password salah."}), 401

@app.route("/api/users", methods=["GET"])
def list_users():
    if not g.get("is_admin", False):
        return jsonify({"error": "Forbidden"}), 403
    return jsonify([{"username": u["username"]} for u in _load_users()])

@app.route("/api/users", methods=["POST"])
def add_user():
    if not g.get("is_admin", False):
        return jsonify({"error": "Forbidden"}), 403
    data = request.get_json(silent=True) or {}
    username = data.get("username", "").strip().lower()
    password = data.get("password", "").strip()
    if not username or not password:
        return jsonify({"error": "Username dan password wajib diisi."}), 400
    if username == _ADMIN_USERNAME.lower():
        return jsonify({"error": "Username sudah digunakan."}), 400
    users = _load_users()
    if any(u["username"].lower() == username for u in users):
        return jsonify({"error": "Username sudah ada."}), 400
    pw_hash = _hash_password(password)
    # Simpan ke Supabase DULU (persistent lintas deploy)
    try:
        from db_services import save_app_user
        ok = save_app_user(username, pw_hash)
        if not ok:
            return jsonify({"error": "Gagal menyimpan akun ke database. Pastikan tabel app_users sudah dibuat di Supabase (lihat supabase_setup.sql)."}), 500
    except Exception as e:
        logger.error(f"[USERS] Gagal simpan user ke Supabase: {e}")
        return jsonify({"error": f"Gagal menyimpan ke Supabase: {e}. Jalankan supabase_setup.sql terlebih dahulu."}), 500
    # Simpan juga ke JSON lokal (cache)
    users.append({"username": username, "password_hash": pw_hash})
    _save_users(users)
    return jsonify({"ok": True, "username": username})

@app.route("/api/users/<username>", methods=["DELETE"])
def delete_user(username):
    if not g.get("is_admin", False):
        return jsonify({"error": "Forbidden"}), 403
    users = _load_users()
    new_users = [u for u in users if u["username"] != username]
    if len(new_users) == len(users):
        return jsonify({"error": "User tidak ditemukan."}), 404
    _save_users(new_users)
    # Hapus dari Supabase
    try:
        from db_services import delete_app_user
        delete_app_user(username)
    except Exception as e:
        logger.warning(f"[USERS] Gagal hapus user dari Supabase: {e}")
    return jsonify({"ok": True})


@app.route("/api/users/<username>/reset-password", methods=["POST"])
def reset_user_password(username):
    if not g.get("is_admin", False):
        return jsonify({"error": "Forbidden"}), 403
    data = request.get_json(silent=True) or {}
    new_password = (data.get("password") or "").strip()
    if not new_password:
        return jsonify({"error": "Password baru wajib diisi."}), 400
    users = _load_users()
    target = next((u for u in users if u["username"] == username), None)
    if not target:
        return jsonify({"error": "User tidak ditemukan."}), 404
    pw_hash = _hash_password(new_password)
    # Update di Supabase dulu
    try:
        from db_services import save_app_user
        ok = save_app_user(username, pw_hash)
        if not ok:
            return jsonify({"error": "Gagal update password di database. Pastikan tabel app_users ada di Supabase."}), 500
    except Exception as e:
        logger.error(f"[USERS] Gagal update password di Supabase: {e}")
        return jsonify({"error": f"Gagal update ke Supabase: {e}"}), 500
    # Update JSON lokal juga
    target["password_hash"] = pw_hash
    _save_users(users)
    return jsonify({"ok": True})


@app.route("/api/users/activity", methods=["GET"])
def get_users_activity():
    if not g.get("is_admin", False):
        return jsonify({"error": "Forbidden"}), 403
    from datetime import datetime, timezone
    import time as _time

    _load_activity()
    users_list = _load_users()

    # Baca push log dari Supabase, fallback ke file lokal
    push_log = []
    try:
        from db_services import fetch_push_logs_from_supabase
        push_log = fetch_push_logs_from_supabase(limit=5000) or []
    except Exception:
        pass
    if not push_log and os.path.exists(PUSH_LOG_FILE):
        try:
            with open(PUSH_LOG_FILE, "r", encoding="utf-8") as f:
                push_log = json.load(f)
        except Exception:
            push_log = []

    push_by_user: dict = {}
    for entry in push_log:
        u = entry.get("username", "unknown")
        if u not in push_by_user:
            push_by_user[u] = {"count": 0, "total_articles": 0, "last_push": None, "last_source": None}
        push_by_user[u]["count"] += 1
        push_by_user[u]["total_articles"] += entry.get("count", 0)
        ts = entry.get("timestamp")
        if ts and (push_by_user[u]["last_push"] is None or ts > push_by_user[u]["last_push"]):
            push_by_user[u]["last_push"] = ts
            push_by_user[u]["last_source"] = entry.get("source")

    def _online_status(last_seen: str | None) -> str:
        if not last_seen:
            return "offline"
        try:
            dt = datetime.fromisoformat(last_seen.replace("Z", "+00:00"))
            if dt.tzinfo is None:
                dt = dt.replace(tzinfo=timezone.utc)
            diff = (_time.time() - dt.timestamp())
            if diff < 300:
                return "online"
            if diff < 1800:
                return "away"
            return "offline"
        except Exception:
            return "offline"

    result = []

    admin_act = _activity_cache.get(_ADMIN_USERNAME, {})
    admin_push = push_by_user.get(_ADMIN_USERNAME, {})
    result.append({
        "username": _ADMIN_USERNAME,
        "is_admin": True,
        "last_login": admin_act.get("last_login"),
        "last_seen": admin_act.get("last_seen"),
        "status": _online_status(admin_act.get("last_seen")),
        "push_count": admin_push.get("count", 0),
        "push_articles": admin_push.get("total_articles", 0),
        "last_push": admin_push.get("last_push"),
        "last_source": admin_push.get("last_source"),
    })

    for u in users_list:
        uname = u["username"]
        act = _activity_cache.get(uname, {})
        push = push_by_user.get(uname, {})
        result.append({
            "username": uname,
            "is_admin": False,
            "last_login": act.get("last_login"),
            "last_seen": act.get("last_seen"),
            "status": _online_status(act.get("last_seen")),
            "push_count": push.get("count", 0),
            "push_articles": push.get("total_articles", 0),
            "last_push": push.get("last_push"),
            "last_source": push.get("last_source"),
        })

    return jsonify(result)


@app.after_request
def add_no_cache_headers(response):
    """Pastikan API responses tidak di-cache oleh browser."""
    if request.path.startswith("/api/") or request.path in ("/settings", "/kb-drafts"):
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
    return response

DATA_DIR = "data"
CONFIG_DIR = "config"
DATA_FILE = os.path.join(DATA_DIR, "scraped_articles.json")
SETTINGS_FILE = os.path.join(DATA_DIR, "settings.json")
SCHEDULER_FILE = os.path.join(CONFIG_DIR, "scheduler_settings.json")
LAST_JOB_FILE = os.path.join(DATA_DIR, "last_job.json")
USERS_FILE = os.path.join(DATA_DIR, "users.json")
USER_ACTIVITY_FILE = os.path.join(DATA_DIR, "user_activity.json")
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(CONFIG_DIR, exist_ok=True)

def _ensure_data_files():
    """Pastikan semua file data wajib ada — buat dengan nilai default jika belum ada."""
    defaults = {
        DATA_FILE: [],
        SETTINGS_FILE: {},
        LAST_JOB_FILE: {},
        USERS_FILE: [],
        USER_ACTIVITY_FILE: {},
    }
    for path, default in defaults.items():
        if not os.path.exists(path):
            try:
                with open(path, "w") as f:
                    json.dump(default, f)
                logger.info(f"[INIT] Buat file baru: {path}")
            except Exception as e:
                logger.warning(f"[INIT] Gagal buat {path}: {e}")

_ensure_data_files()

_activity_cache: dict = {}
_activity_dirty: dict = {}

def _load_activity() -> dict:
    global _activity_cache
    try:
        if os.path.exists(USER_ACTIVITY_FILE):
            with open(USER_ACTIVITY_FILE, "r", encoding="utf-8") as f:
                _activity_cache = json.load(f)
    except Exception:
        _activity_cache = {}
    # Sinkronisasi dari Supabase — override data JSON lokal
    try:
        from db_services import fetch_user_activity
        remote = fetch_user_activity()
        if remote:
            for uname, row in remote.items():
                if uname not in _activity_cache:
                    _activity_cache[uname] = {}
                if row.get("last_login"):
                    _activity_cache[uname]["last_login"] = row["last_login"]
                # last_seen dari Supabase hanya dipakai jika lebih baru dari lokal
                remote_ls = row.get("last_seen")
                local_ls = _activity_cache[uname].get("last_seen")
                if remote_ls and (not local_ls or remote_ls > local_ls):
                    _activity_cache[uname]["last_seen"] = remote_ls
    except Exception as e:
        logger.warning(f"[ACTIVITY] Gagal sinkron dari Supabase: {e}")
    return _activity_cache

def _save_activity() -> None:
    os.makedirs(DATA_DIR, exist_ok=True)
    with open(USER_ACTIVITY_FILE, "w", encoding="utf-8") as f:
        json.dump(_activity_cache, f, indent=2)

def _update_last_seen(username: str) -> None:
    import time as _time
    now_ts = _time.time()
    prev = _activity_dirty.get(username, 0)
    if now_ts - prev < 30:
        return
    _activity_dirty[username] = now_ts
    if not _activity_cache:
        _load_activity()
    if username not in _activity_cache:
        _activity_cache[username] = {}
    _activity_cache[username]["last_seen"] = _now_iso()
    _save_activity()

def _record_login(username: str) -> None:
    if not _activity_cache:
        _load_activity()
    if username not in _activity_cache:
        _activity_cache[username] = {}
    now = _now_iso()
    _activity_cache[username]["last_login"] = now
    _save_activity()
    # Persist login ke Supabase
    try:
        from db_services import save_user_activity
        save_user_activity(username, last_login=now)
    except Exception as e:
        logger.warning(f"[ACTIVITY] Gagal simpan login ke Supabase: {e}")

_load_activity()

DEFAULT_SCHEDULER = {
    "enabled": False,
    "interval": "manual",   # "manual" | "daily" | "weekly"
    "day_of_week": "mon",   # used when weekly (mon/tue/wed/thu/fri/sat/sun)
    "time_of_day": "06:00", # HH:MM
    "url": "",
    "scrape_mode": "full",  # "full" | "kb"
    "incremental": True,
    "rss_sources": [],       # list of {url, label, max_items}
    "last_run_at": None,
    "last_run_articles_added": 0,
    "last_run_url": "",
    "last_run_mode": "full",
}

# APScheduler — background, survives across requests
_scheduler = BackgroundScheduler(timezone="Asia/Jakarta")
_scheduler.start()


def _load_scheduler_settings() -> dict:
    if os.path.exists(SCHEDULER_FILE):
        try:
            with open(SCHEDULER_FILE, "r", encoding="utf-8") as f:
                saved = json.load(f)
                return {**DEFAULT_SCHEDULER, **saved}
        except Exception:
            pass
    return dict(DEFAULT_SCHEDULER)


def _save_scheduler_settings(data: dict):
    with open(SCHEDULER_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _next_run_iso() -> str | None:
    """Kembalikan waktu eksekusi berikutnya dari APScheduler sebagai ISO string."""
    job = _scheduler.get_job("scheduled_scrape")
    if job and job.next_run_time:
        return job.next_run_time.strftime("%Y-%m-%dT%H:%M:%S")
    return None


def _run_scheduled_rss(rss_sources: list) -> int:
    """Jalankan scraping RSS untuk semua sumber terjadwal. Kembalikan jumlah artikel baru."""
    total_added = 0
    try:
        import feedparser as _fp
        from bs4 import BeautifulSoup as _BS
        import re as _re
        from datetime import datetime as _dt
        import time as _t
    except ImportError as e:
        logging.warning(f"[SCHEDULER-RSS] Library tidak tersedia: {e}")
        return 0

    for src in rss_sources:
        url = (src.get("url") or "").strip()
        label = src.get("label") or url
        max_items = int(src.get("max_items") or 10)
        if not url:
            continue
        try:
            feed = _fp.parse(url)
            entries = feed.get("entries", [])[:max_items]
            for entry in entries:
                title = entry.get("title", "").strip() or "(Tanpa Judul)"
                link = entry.get("link", "")
                content_raw = ""
                if entry.get("content"):
                    content_raw = entry["content"][0].get("value", "")
                if not content_raw:
                    content_raw = entry.get("summary", "") or entry.get("description", "")
                try:
                    content = _BS(content_raw, "html.parser").get_text(separator="\n").strip()
                except Exception:
                    content = _re.sub(r"<[^>]+>", " ", content_raw).strip()
                if not content:
                    content = title
                pub_date = ""
                if entry.get("published_parsed"):
                    try:
                        pub_date = _dt(*entry["published_parsed"][:6]).strftime("%Y-%m-%dT%H:%M:%S")
                    except Exception:
                        pass
                _make_kb_draft(title, content, link, pub_date, source_tag="rss")
                total_added += 1
            logging.info(f"[SCHEDULER-RSS] {label}: {len(entries)} artikel diproses.")
        except Exception as e:
            logging.warning(f"[SCHEDULER-RSS] Gagal fetch {label}: {e}")
    return total_added


def _run_scheduled_scrape():
    """Dipanggil oleh APScheduler. Gunakan settings yang tersimpan."""
    cfg = _load_scheduler_settings()
    url = cfg.get("url", "").strip()
    mode = cfg.get("scrape_mode", "full")
    incremental = cfg.get("incremental", True)
    rss_sources = cfg.get("rss_sources", [])

    has_web = bool(url)
    has_rss = bool(rss_sources)

    if not has_web and not has_rss:
        logging.warning("[SCHEDULER] Tidak ada sumber yang dikonfigurasi, scraping dibatalkan.")
        return

    with state_lock:
        if scrape_state["running"]:
            logging.warning("[SCHEDULER] Scraping sedang berjalan, jadwal dilewati.")
            return

    added = 0

    # ── Web scraping ──
    if has_web:
        logging.info(f"[SCHEDULER] Web scraping terjadwal: {url} | mode={mode}")
        settings = _load_settings()
        before_count = len(_load_articles())
        _run_scrape(url, settings, mode, scheduled=True, incremental=incremental)
        after_count = len(_load_articles())
        added += max(0, after_count - before_count)

    # ── RSS scraping ──
    if has_rss:
        logging.info(f"[SCHEDULER] RSS scraping terjadwal: {len(rss_sources)} sumber")
        added += _run_scheduled_rss(rss_sources)

    cfg = _load_scheduler_settings()
    cfg["last_run_at"] = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")
    cfg["last_run_articles_added"] = added
    cfg["last_run_url"] = url
    cfg["last_run_mode"] = mode
    _save_scheduler_settings(cfg)
    logging.info(f"[SCHEDULER] Selesai. {added} artikel baru ditambahkan.")


def _apply_scheduler(cfg: dict):
    """Terapkan job APScheduler sesuai settings. Hapus job lama dahulu."""
    _scheduler.remove_job("scheduled_scrape") if _scheduler.get_job("scheduled_scrape") else None

    if not cfg.get("enabled") or cfg.get("interval") == "manual":
        return  # Tidak ada jadwal

    time_str = cfg.get("time_of_day", "06:00")
    try:
        h, m = time_str.split(":")
        hour, minute = int(h), int(m)
    except Exception:
        hour, minute = 6, 0

    interval = cfg.get("interval")
    if interval == "daily":
        trigger = CronTrigger(hour=hour, minute=minute, timezone="Asia/Jakarta")
    elif interval == "weekly":
        dow = cfg.get("day_of_week", "mon")
        trigger = CronTrigger(day_of_week=dow, hour=hour, minute=minute, timezone="Asia/Jakarta")
    else:
        return

    _scheduler.add_job(
        _run_scheduled_scrape,
        trigger=trigger,
        id="scheduled_scrape",
        replace_existing=True,
        misfire_grace_time=300,
    )
    logging.info(f"[SCHEDULER] Job terdaftar: interval={interval}, jam={hour:02d}:{minute:02d}")


# Terapkan scheduler dari settings yang tersimpan saat startup
_apply_scheduler(_load_scheduler_settings())

DEFAULT_SETTINGS = {
    "article_link_selector": 'a[href*="/berita/"]',
    "next_page_selector": 'a[rel="next"], a.next, .pagination a',
    "title_selector": "h1, h2, .title, .news-title, .post-title",
    "date_selector": ".date, .news-date, time, .published-date, .post-date",
    "content_selector": ".ck-content, .post-content, .news-content, .article-content, .content, article, .entry-content",
}


def _load_settings() -> dict:
    if os.path.exists(SETTINGS_FILE):
        try:
            with open(SETTINGS_FILE, "r", encoding="utf-8") as f:
                saved = json.load(f)
                merged = {**DEFAULT_SETTINGS, **saved}
                return merged
        except Exception:
            pass
    return dict(DEFAULT_SETTINGS)


def _save_settings(data: dict):
    with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _load_last_job() -> dict | None:
    """Load parameter scrape terakhir yang dijalankan."""
    if os.path.exists(LAST_JOB_FILE):
        try:
            with open(LAST_JOB_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return None


def _save_last_job(url: str, mode: str, start_date: date | None, end_date: date | None):
    """Simpan parameter scrape terakhir ke disk."""
    job = {
        "url": url,
        "mode": mode,
        "start_date": start_date.isoformat() if start_date else None,
        "end_date": end_date.isoformat() if end_date else None,
        "saved_at": datetime.now().isoformat(),
    }
    with open(LAST_JOB_FILE, "w", encoding="utf-8") as f:
        json.dump(job, f, ensure_ascii=False, indent=2)


# State global untuk progress tracking
scrape_state = {
    "running": False,
    "cancelled": False,   # set by reset-all to abort in-flight scrape saves
    "phase": "idle",      # idle, listing, scraping, done
    "current": 0,
    "total": 0,
    "success": 0,
    "partial": 0,
    "failed": 0,
    "duplicate": 0,
    "logs": [],
    "articles": [],
}
state_lock = threading.Lock()


def _load_articles() -> list:
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return []


def _save_articles(data: list):
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _progress_callback(msg, **kwargs):
    with state_lock:
        scrape_state["logs"].append(msg)
        for k, v in kwargs.items():
            if k in scrape_state:
                scrape_state[k] = v


def _run_scrape(url: str, settings: dict, mode: str,
                start_date: date | None = None, end_date: date | None = None,
                incremental: bool = True, scheduled: bool = False):
    run_label = "[SCHEDULED]" if scheduled else "[MANUAL]"
    # Simpan parameter job ini agar bisa di-restart setelah reset
    if not scheduled:
        _save_last_job(url, mode, start_date, end_date)
    with state_lock:
        scrape_state["running"] = True
        scrape_state["cancelled"] = False   # clear any previous cancellation
        scrape_state["phase"] = "listing"
        scrape_state["current"] = 0
        scrape_state["total"] = 0
        scrape_state["success"] = 0
        scrape_state["partial"] = 0
        scrape_state["failed"] = 0
        scrape_state["duplicate"] = 0
        scrape_state["logs"] = [f"{run_label} Memulai scraping: {url}"]
        scrape_state["articles"] = []

    try:
        # ── Auto-detect selectors untuk non-kemlu sites ───────────────────────
        if not is_kemlu_url(url):
            detected = auto_detect_selectors(url, log_fn=_progress_callback)
            if detected:
                settings = {**settings, **detected}
                # Simpan selector yang terdeteksi ke file settings agar persistent
                try:
                    cfg = _load_settings()
                    cfg.update(detected)
                    _save_settings(cfg)
                except Exception:
                    pass

        # Incremental: load existing untuk deduplication; Full refresh: start clean
        existing = _load_articles() if incremental else []
        if not incremental:
            _progress_callback(f"{run_label} Mode full refresh — data lama dibersihkan.")

        # Merged list yang akan di-update secara inkremental
        merged = list(existing)
        merged_urls = {a["url"] for a in existing}

        def _article_callback(article):
            """Dipanggil setelah setiap artikel berhasil di-scrape — simpan segera ke disk."""
            url_key = article.get("url", "")
            with state_lock:
                # If reset was called while we were scraping, discard this write
                if scrape_state.get("cancelled"):
                    return
                if url_key and url_key in merged_urls:
                    return
                if url_key:
                    merged_urls.add(url_key)
                merged.append(article)
                # Update running counters
                s = article.get("status", "")
                if s == "success":
                    scrape_state["success"] += 1
                elif s == "partial":
                    scrape_state["partial"] += 1
                elif s == "failed":
                    scrape_state["failed"] += 1
            _save_articles(merged)

        scrape_all(
            url,
            settings=settings,
            mode=mode,
            existing_articles=existing,
            progress_callback=_progress_callback,
            start_date=start_date,
            end_date=end_date,
            article_callback=_article_callback,
        )

        # Final save — skip if reset was called mid-scrape
        with state_lock:
            was_cancelled = scrape_state.get("cancelled", False)
        if not was_cancelled:
            _save_articles(merged)

        with state_lock:
            scrape_state["articles"] = merged
            scrape_state["phase"] = "done"
    except Exception as e:
        with state_lock:
            scrape_state["logs"].append(f"FATAL ERROR: {str(e)}")
            scrape_state["phase"] = "done"
    finally:
        with state_lock:
            scrape_state["running"] = False


# ─── Routes ─────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/article/<article_id>")
def article_detail(article_id):
    articles = _load_articles()
    article = next((a for a in articles if a["id"] == article_id), None)
    return render_template("article.html", article=article)


@app.route("/settings", methods=["GET"])
def get_settings():
    return jsonify(_load_settings())


@app.route("/settings", methods=["POST"])
def post_settings():
    data = request.get_json(force=True)
    allowed_keys = {
        "article_link_selector", "next_page_selector",
        "title_selector", "date_selector", "content_selector",
    }
    current = _load_settings()
    for k in allowed_keys:
        if k in data:
            current[k] = str(data[k]).strip()
    _save_settings(current)
    return jsonify({"status": "ok", "settings": current})


@app.route("/api/detect-selectors", methods=["POST"])
def api_detect_selectors():
    """
    Deteksi otomatis CSS selector dari URL yang diberikan.
    Body: { "url": "https://..." }
    """
    data = request.get_json(force=True)
    url = (data.get("url") or "").strip()
    if not url:
        return jsonify({"error": "URL diperlukan"}), 400
    logs: list[str] = []
    detected = auto_detect_selectors(url, log_fn=lambda m: logs.append(m))
    if detected:
        current = _load_settings()
        current.update(detected)
        _save_settings(current)
    return jsonify({"detected": detected, "logs": logs})


def _parse_date_param(s: str | None) -> date | None:
    """Parse 'yyyy-mm-dd' string ke datetime.date, return None jika gagal."""
    if not s:
        return None
    try:
        return datetime.strptime(s.strip(), "%Y-%m-%d").date()
    except ValueError:
        return None


@app.route("/api/scrape", methods=["POST"])
def api_scrape():
    with state_lock:
        if scrape_state["running"]:
            return jsonify({"error": "Scraping sedang berjalan"}), 409
    data = request.get_json(force=True)
    url = (data.get("url") or "").strip()
    mode = (data.get("mode") or "full").strip()
    if mode not in ("list", "full", "kb"):
        mode = "full"
    if not url:
        return jsonify({"error": "URL tidak boleh kosong"}), 400
    if not url.startswith("http"):
        return jsonify({"error": "URL tidak valid, harus dimulai dengan http:// atau https://"}), 400

    # ── Date range filter ──
    date_filter = (data.get("date_filter") or "all").strip()
    today = date.today()
    start_date: date | None = None
    end_date: date | None = today  # cap upper bound at today

    if date_filter == "last_7":
        start_date = today - timedelta(days=7)
    elif date_filter == "last_30":
        start_date = today - timedelta(days=30)
    elif date_filter == "custom":
        start_date = _parse_date_param(data.get("start_date"))
        end_date = _parse_date_param(data.get("end_date")) or today
    else:
        end_date = None  # "all" → no cap

    settings = _load_settings()
    threading.Thread(
        target=_run_scrape,
        args=(url, settings, mode, start_date, end_date),
        daemon=True,
    ).start()
    return jsonify({"status": "started", "date_filter": date_filter,
                    "start_date": start_date.isoformat() if start_date else None,
                    "end_date": end_date.isoformat() if end_date else None})


@app.route("/api/progress")
def api_progress():
    with state_lock:
        return jsonify({
            "running": scrape_state["running"],
            "phase": scrape_state["phase"],
            "current": scrape_state["current"],
            "total": scrape_state["total"],
            "success": scrape_state["success"],
            "partial": scrape_state["partial"],
            "failed": scrape_state["failed"],
            "duplicate": scrape_state["duplicate"],
            "logs": scrape_state["logs"][-50:],  # kirim 50 log terakhir
        })


@app.route("/api/articles")
def api_articles():
    articles = _load_articles()
    return jsonify(articles)


@app.route("/api/article/<article_id>")
def api_article(article_id):
    articles = _load_articles()
    article = next((a for a in articles if a["id"] == article_id), None)
    if not article:
        return jsonify({"error": "Artikel tidak ditemukan"}), 404
    return jsonify(article)


@app.route("/api/article/<article_id>/format", methods=["POST"])
def api_format_article(article_id):
    """Rapikan konten artikel menggunakan AI."""
    from ai_services import get_openai_client, check_openai_available
    if not check_openai_available():
        return jsonify({"error": "OpenAI API key tidak ditemukan. Fitur ini membutuhkan OPENAI_API_KEY."}), 503

    articles = _load_articles()
    article = next((a for a in articles if a["id"] == article_id), None)
    if not article:
        return jsonify({"error": "Artikel tidak ditemukan"}), 404

    data = request.get_json(force=True) or {}
    save = data.get("save", False)

    title = article.get("title", "")
    content = article.get("content", "")
    if not content:
        return jsonify({"error": "Artikel tidak memiliki konten."}), 400

    try:
        client = get_openai_client()
        prompt = (
            "Kamu adalah editor konten berita profesional. Tugasmu adalah mengekstrak dan menyajikan HANYA informasi penting dari artikel berikut dalam format Markdown yang rapi dan presisi.\n\n"
            f"Judul artikel: {title}\n\n"
            f"Konten asli:\n{content[:6000]}\n\n"
            "INSTRUKSI KETAT:\n\n"
            "**Yang HARUS dibuang (jangan masukkan sama sekali):**\n"
            "- Iklan, promo, ajakan subscribe/follow\n"
            "- Navigasi website, footer, cookie notice, disclaimer boilerplate\n"
            "- Kalimat basa-basi, pembuka/penutup tidak informatif\n"
            "- Informasi duplikat atau pengulangan\n"
            "- Opini wartawan yang tidak didukung fakta\n"
            "- Informasi yang tidak relevan dengan topik utama artikel\n\n"
            "**Yang HARUS dipertahankan:**\n"
            "- Fakta utama: apa, siapa, kapan, di mana, mengapa, bagaimana\n"
            "- Angka, statistik, tanggal, nama resmi\n"
            "- Kutipan langsung yang penting\n"
            "- Konteks dan latar belakang yang relevan\n\n"
            "**Format output (Markdown):**\n"
            "- Gunakan `##` untuk sub-topik jika ada lebih dari satu topik\n"
            "- Gunakan bullet points (`-`) untuk daftar fakta atau poin-poin\n"
            "- Gunakan **bold** untuk nama, jabatan, angka, atau istilah kunci\n"
            "- Gunakan paragraf prose untuk narasi yang mengalir\n"
            "- Pisahkan bagian dengan satu baris kosong\n"
            "- Gunakan bahasa yang sama dengan artikel asli (jangan terjemahkan)\n\n"
            "PENTING: Tulis HANYA konten Markdown. Jangan tambahkan kata pengantar, penutup, atau komentar apapun dari kamu."
        )
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2500,
            temperature=0.1,
        )
        formatted = response.choices[0].message.content.strip()

        if save:
            for a in articles:
                if a["id"] == article_id:
                    a["content"] = formatted
                    a["formatted_by_ai"] = True
                    break
            _save_articles(articles)

        return jsonify({"status": "ok", "formatted_content": formatted})
    except Exception as e:
        logger.error(f"[FORMAT] Error: {e}")
        return jsonify({"error": str(e)}), 500


def _detect_arabic_ratio(text: str) -> float:
    """Hitung persentase karakter Arab dalam teks."""
    if not text:
        return 0.0
    arab_count = sum(1 for c in text if '\u0600' <= c <= '\u06FF' or '\u0750' <= c <= '\u077F' or '\uFB50' <= c <= '\uFDFF' or '\uFE70' <= c <= '\uFEFF')
    letter_count = sum(1 for c in text if c.isalpha())
    return arab_count / max(letter_count, 1)

_ARABIC_RECONSTRUCTION_SYSTEM = (
    "Kamu adalah ahli rekonstruksi teks Arab dari hasil OCR dan spesialis bahasa Arab klasik/modern.\n\n"
    "KONTEKS PENTING:\n"
    "Teks yang diberikan adalah hasil OCR dari kitab/dokumen berbahasa Arab. OCR sering menghasilkan:\n"
    "- Huruf Arab yang tercerai-berai (misal: 'ا ل ح م د' → seharusnya 'الحمد')\n"
    "- Spasi yang salah posisi di tengah kata\n"
    "- Karakter aneh atau simbol pengganti huruf yang tidak terbaca\n"
    "- Urutan kata yang terbalik atau baris yang tertukar\n"
    "- Harakat/syakal yang terpisah dari huruf aslinya\n\n"
    "TUGAS UTAMAMU:\n"
    "1. REKONSTRUKSI setiap kata Arab ke bentuk yang benar — sambungkan huruf-huruf yang tercerai\n"
    "2. PERBAIKI spasi yang salah di dalam kata (spasi hanya antar kata, bukan di tengah kata)\n"
    "3. PERTAHANKAN semua harakat/syakal (ـَ ـِ ـُ ـّ dll) jika ada, letakkan di posisi benar\n"
    "4. PERTAHANKAN semua konten — jangan tambah atau hapus makna\n"
    "5. PERTAHANKAN bahasa asli — jangan terjemahkan ke Indonesia\n"
    "6. Untuk teks campuran Arab-Indonesia: rekonstruksi Arab, pertahankan Indonesia apa adanya\n\n"
    "STANDAR KUALITAS OUTPUT:\n"
    "- Kata Arab harus terbaca seperti kitab yang dicetak normal\n"
    "- Gunakan Unicode Arab yang benar (bukan karakter terpisah)\n"
    "- Jangan gunakan harakat jika tidak ada di teks asli\n"
)

# ── System message untuk SEMUA mode (non-Arab) ──────────────────────────────
_AINA_SYSTEM = (
    "Kamu adalah kurator Knowledge Base untuk AINA — asisten AI Islam berbahasa Indonesia dan Arab.\n\n"
    "KONTEKS AINA:\n"
    "AINA menjawab pertanyaan seputar fikih Islam, kebijakan keagamaan Indonesia, berita dunia Muslim, "
    "fatwa ulama, teks kitab Arab, dan isu sosial-politik di negara Muslim. "
    "Output yang kamu buat langsung masuk ke database pengetahuan AINA dan digunakan sebagai sumber jawaban.\n\n"
    "PRINSIP UTAMA:\n"
    "1. EKSTRAK semua informasi yang berguna untuk menjawab pertanyaan — jangan potong fakta penting\n"
    "2. HAPUS noise: iklan, navigasi, promo, basa-basi, disclaimer boilerplate, duplikasi, opini tanpa fakta\n"
    "3. PERTAHANKAN: nama tokoh, jabatan, angka, tanggal, kutipan, istilah teknis (Arab/Latin), sumber\n"
    "4. TABEL: jika ada data tabel, WAJIB format sebagai tabel Markdown yang rapi:\n"
    "   | Kolom 1 | Kolom 2 | Kolom 3 |\n"
    "   |---------|---------|----------|\n"
    "   | Data    | Data    | Data     |\n"
    "   Pastikan alignment kolom konsisten dan semua sel terisi.\n"
    "5. BAHASA: gunakan bahasa yang sama dengan teks sumber (jangan terjemahkan)\n"
    "6. OUTPUT: HANYA Markdown bersih — tanpa kata pengantar, tanpa penutup, tanpa komentar dari kamu\n"
)

_FORMAT_SYSTEM = {
    "berita": (
        "Setelah rekonstruksi Arab selesai, sajikan hasilnya sebagai konten berita KB-AINA yang rapi.\n"
        "- Ekstrak fakta penting: apa, siapa, kapan, di mana, mengapa, bagaimana\n"
        "- Format: `##` sub-topik, `-` fakta, **bold** nama/jabatan/angka kunci\n"
        "- Jika ada data tabel dalam teks, sajikan sebagai tabel Markdown\n"
        "- Pertahankan semua teks Arab yang sudah direkonstruksi\n"
        "Output: HANYA Markdown. Tanpa pengantar atau komentar."
    ),
    "kitab": (
        "Setelah rekonstruksi Arab selesai, strukturkan sebagai teks kitab KB-AINA.\n"
        "- Gunakan `##` untuk Bab/Fasal/Pasal yang terdeteksi\n"
        "- Letakkan teks Arab dalam blok tersendiri, syarah/terjemahan sesudahnya\n"
        "- Pertahankan nomor poin/masalah/fasal, istilah fikih dan ushul\n"
        "- Jika ada tabel (misalnya perbandingan madzhab), gunakan tabel Markdown\n"
        "Output: HANYA teks kitab terstruktur dalam Markdown. Tanpa komentar."
    ),
    "laporan": (
        "Setelah rekonstruksi Arab selesai, susun sebagai laporan formal KB-AINA.\n"
        "- Mulai dengan **Ringkasan Eksekutif** (2-3 kalimat inti)\n"
        "- Gunakan `##` untuk: Latar Belakang, Temuan Utama, Analisis, Rekomendasi\n"
        "- Jika ada data angka/statistik: sajikan sebagai tabel Markdown\n"
        "- Pertahankan data, angka, nama resmi, teks Arab penting\n"
        "Output: HANYA konten laporan Markdown. Tanpa pengantar."
    ),
    "ringkasan": (
        "Setelah rekonstruksi Arab selesai, buat ringkasan singkat untuk KB-AINA.\n"
        "- Satu kalimat konteks di atas, lalu 3-7 poin bullet list `-` paling penting\n"
        "- Sertakan teks Arab kunci, nama tokoh, tanggal, angka yang relevan\n"
        "- Setiap poin harus mandiri dan cukup untuk menjawab pertanyaan\n"
        "Output: HANYA ringkasan Markdown. Tanpa kata pengantar."
    ),
    "poin": (
        "Setelah rekonstruksi Arab selesai, ubah menjadi daftar poin informatif KB-AINA.\n"
        "- Setiap fakta/hukum/masalah = 1 bullet point `-` yang berdiri sendiri\n"
        "- **bold** di awal setiap poin untuk kata kunci (hukum, nama, istilah)\n"
        "- Kelompokkan dengan `##` jika ada kategori berbeda\n"
        "- Pertahankan istilah Arab yang penting, jangan terjemahkan\n"
        "Output: HANYA daftar poin Markdown. Tanpa narasi pembuka."
    ),
    "briefing": (
        "Setelah rekonstruksi Arab selesai, buat briefing intelijen/diplomatik untuk KB-AINA.\n"
        "## SITUASI\n(1-2 kalimat gambaran keseluruhan)\n"
        "## FAKTA KUNCI\n(bullet list: angka, nama, tanggal terverifikasi)\n"
        "## AKTOR\n(siapa yang terlibat, peran masing-masing)\n"
        "## IMPLIKASI\n(apa artinya untuk kebijakan/umat/hubungan antar negara)\n"
        "- Jika ada data komparatif, gunakan tabel Markdown\n"
        "Output: HANYA konten briefing Markdown. Tanpa pengantar."
    ),
}

# ── Prompts untuk mode non-Arab (dengan system message _AINA_SYSTEM) ────────
FORMAT_PROMPTS = {
    "berita": lambda title, content: (
        "Kamu adalah kurator KB-AINA. Proses teks berikut menjadi konten berita yang siap digunakan AINA untuk menjawab pertanyaan.\n\n"
        + (f"**Judul:** {title}\n\n" if title else "")
        + f"**Teks sumber:**\n{content}\n\n"
        "## INSTRUKSI KETAT\n\n"
        "**HAPUS (jangan masukkan sama sekali):**\n"
        "- Iklan, promo, banner, ajakan subscribe/follow/like\n"
        "- Menu navigasi, breadcrumb, footer, copyright, cookie notice\n"
        "- Kalimat pembuka/penutup basa-basi ('Baca juga...', 'Artikel terkait...', dll)\n"
        "- Opini penulis tanpa dasar fakta, spekulasi\n"
        "- Konten duplikat atau pengulangan informasi yang sama\n\n"
        "**WAJIB PERTAHANKAN:**\n"
        "- Fakta 5W+1H (apa, siapa, kapan, di mana, mengapa, bagaimana)\n"
        "- Angka, statistik, persentase, nilai, jumlah\n"
        "- Tanggal dan waktu spesifik\n"
        "- Nama lengkap tokoh beserta jabatan/gelar\n"
        "- Kutipan langsung yang substantif (pakai > blockquote)\n"
        "- Nama lembaga, kebijakan, regulasi, produk hukum\n"
        "- Konteks sejarah/latar belakang yang relevan\n\n"
        "**FORMAT OUTPUT (Markdown):**\n"
        "- `##` untuk sub-topik utama jika ada beberapa aspek berbeda\n"
        "- `-` bullet untuk daftar fakta\n"
        "- **bold** untuk nama, jabatan, angka, istilah kunci\n"
        "- `> kutipan` untuk kutipan langsung penting\n"
        "- TABEL: jika ada data perbandingan/statistik tabel, wajib gunakan format:\n"
        "  | Header | Header |\n  |--------|--------|\n  | Data   | Data   |\n"
        "- Paragraf prose untuk narasi yang membutuhkan alur\n"
        "- Gunakan bahasa yang sama dengan sumber\n\n"
        "Output: HANYA konten Markdown. Tanpa kata pengantar, penutup, atau komentar."
    ),
    "kitab": lambda title, content: (
        "Kamu adalah kurator KB-AINA spesialis teks keagamaan. Rapikan dan strukturkan teks kitab/agama berikut agar optimal untuk retrieval AINA.\n\n"
        + (f"**Kitab/Judul:** {title}\n\n" if title else "")
        + f"**Teks sumber:**\n{content}\n\n"
        "## INSTRUKSI\n\n"
        "**Rekonstruksi & Pembersihan:**\n"
        "- Perbaiki OCR: sambungkan huruf Arab yang terpisah, perbaiki spasi yang salah\n"
        "- Buang karakter noise, watermark, header/footer halaman yang tidak relevan\n"
        "- PERTAHANKAN semua konten — jangan hapus atau tambah makna apapun\n\n"
        "**Struktur:**\n"
        "- `##` untuk Bab / Fasal / Pasal / Kitab yang terdeteksi\n"
        "- `###` untuk sub-bagian (masalah, kaidah, dll)\n"
        "- Teks Arab dalam blok terpisah, syarah/terjemahan/komentar sesudahnya\n"
        "- Pertahankan penomoran (1. 2. 3. atau أولاً ثانياً dll)\n"
        "- Istilah fikih/ushul/hadis: pertahankan dalam Arab, tambah transliterasi jika ada\n\n"
        "**Tabel:** jika ada perbandingan madzhab, hukum, atau data terstruktur:\n"
        "| Aspek | Hanafi | Maliki | Syafi'i | Hanbali |\n"
        "|-------|--------|--------|---------|--------|\n"
        "| ...   | ...    | ...    | ...     | ...    |\n\n"
        "Output: HANYA teks kitab terstruktur dalam Markdown. Tanpa komentar dari kamu."
    ),
    "laporan": lambda title, content: (
        "Kamu adalah kurator KB-AINA spesialis laporan formal. Susun ulang teks berikut menjadi laporan terstruktur yang informatif untuk AINA.\n\n"
        + (f"**Judul Laporan:** {title}\n\n" if title else "")
        + f"**Konten sumber:**\n{content}\n\n"
        "## INSTRUKSI\n\n"
        "**Struktur laporan (gunakan `##` untuk setiap bagian):**\n"
        "1. **Ringkasan Eksekutif** — 2-4 kalimat inti yang menjawab: apa yang terjadi, mengapa penting\n"
        "2. **Latar Belakang** — konteks yang dibutuhkan untuk memahami laporan\n"
        "3. **Temuan / Fakta Utama** — fakta-fakta kunci dengan angka, nama, tanggal\n"
        "4. **Analisis** — implikasi, sebab-akibat, signifikansi (hanya jika ada dalam sumber)\n"
        "5. **Rekomendasi / Tindak Lanjut** — (hanya jika ada dalam sumber)\n\n"
        "**Data & Tabel:**\n"
        "- Semua data statistik, angka perbandingan, atau data tabel WAJIB disajikan sebagai tabel Markdown\n"
        "- Format tabel: rata kiri untuk teks, rata kanan untuk angka\n"
        "- Contoh:\n"
        "  | Indikator | Tahun 2023 | Tahun 2024 | Perubahan |\n"
        "  |-----------|-----------|------------|----------|\n"
        "  | ...       | ...       | ...        | ...      |\n\n"
        "**Gaya:** formal, objektif, **bold** untuk istilah kunci dan poin krusial\n\n"
        "Output: HANYA konten laporan Markdown. Tanpa pengantar atau komentar dari kamu."
    ),
    "ringkasan": lambda title, content: (
        "Kamu adalah kurator KB-AINA. Buat ringkasan yang padat dan informatif — cukup untuk AINA menjawab pertanyaan tentang topik ini.\n\n"
        + (f"**Topik:** {title}\n\n" if title else "")
        + f"**Teks sumber:**\n{content}\n\n"
        "## INSTRUKSI\n\n"
        "**Struktur output:**\n"
        "- Satu kalimat pembuka yang merangkum inti topik (tanpa header)\n"
        "- Bullet list `-` berisi 4-8 poin PALING PENTING, masing-masing 1-2 kalimat\n"
        "- Setiap poin harus mandiri (tidak merujuk 'di atas' atau 'hal tersebut')\n"
        "- Urutan: dari yang paling signifikan ke yang kurang signifikan\n\n"
        "**Wajib masukkan jika ada:**\n"
        "- Nama tokoh + jabatan\n"
        "- Tanggal/periode spesifik\n"
        "- Angka dan statistik kunci\n"
        "- Keputusan/fatwa/kebijakan yang dihasilkan\n"
        "- Nama lembaga atau regulasi\n\n"
        "**Tabel:** jika ada data yang lebih jelas disajikan sebagai tabel, gunakan tabel Markdown\n\n"
        "Output: HANYA ringkasan Markdown. Tanpa kata 'Ringkasan:', pengantar, atau komentar."
    ),
    "poin": lambda title, content: (
        "Kamu adalah kurator KB-AINA. Ubah teks berikut menjadi daftar poin informatif yang optimal untuk retrieval Q&A.\n\n"
        + (f"**Topik:** {title}\n\n" if title else "")
        + f"**Teks sumber:**\n{content}\n\n"
        "## INSTRUKSI\n\n"
        "**Aturan poin:**\n"
        "- Setiap fakta/hukum/keputusan/data = 1 poin bullet `-` yang berdiri sendiri\n"
        "- Poin harus lengkap: bisa dibaca dan dimengerti tanpa membaca poin lain\n"
        "- **Bold** di awal setiap poin untuk kata kunci (nama hukum, tokoh, istilah, lembaga)\n"
        "- Satu poin = satu informasi — jangan gabungkan dua fakta berbeda\n"
        "- Minimal 6 poin, maksimal sesuai kedalaman konten sumber\n\n"
        "**Pengelompokan:**\n"
        "- Gunakan `##` untuk kategori berbeda jika konten memiliki beberapa aspek\n"
        "- Urutan: hukum/keputusan utama → detail → konteks → pengecualian\n\n"
        "**Buang:** iklan, basa-basi, navigasi, konten tidak informatif\n\n"
        "**Tabel:** jika ada data perbandingan atau matriks, gunakan tabel Markdown setelah poin terkait\n\n"
        "Output: HANYA daftar poin Markdown. Tanpa narasi pembuka atau penutup."
    ),
    "briefing": lambda title, content: (
        "Kamu adalah analis KB-AINA. Buat briefing intelijen/diplomatik yang presisi dari teks berikut untuk digunakan AINA menjawab pertanyaan kebijakan dan geopolitik Islam.\n\n"
        + (f"**Subjek:** {title}\n\n" if title else "")
        + f"**Sumber:**\n{content}\n\n"
        "## INSTRUKSI — Format briefing wajib:\n\n"
        "## SITUASI\n"
        "Satu paragraf singkat (2-3 kalimat): apa yang terjadi, di mana, kapan, mengapa signifikan.\n\n"
        "## FAKTA KUNCI\n"
        "Bullet list fakta terverifikasi dari teks:\n"
        "- **[Label]**: deskripsi singkat dengan angka/nama/tanggal spesifik\n"
        "- Hanya fakta yang ADA di teks sumber — jangan asumsi atau tambah\n\n"
        "## AKTOR\n"
        "- **[Nama/Lembaga]** — peran/posisi/tindakan mereka\n"
        "(Jika banyak aktor, gunakan tabel: | Aktor | Posisi | Tindakan |)\n\n"
        "## IMPLIKASI\n"
        "- Dampak kebijakan, hukum, atau sosial yang disebutkan dalam teks\n"
        "- Apa yang perlu diketahui AINA untuk menjawab pertanyaan terkait\n\n"
        "**Gaya:** singkat, presisi, faktual — bahasa netral seperti laporan intelijen\n\n"
        "Output: HANYA konten briefing Markdown. Tanpa pengantar atau komentar dari kamu."
    ),
}

@app.route("/api/format-text", methods=["POST"])
def api_format_text():
    """Rapikan teks artikel. Auto-deteksi Arab untuk rekonstruksi OCR yang kuat."""
    from ai_services import get_openai_client, check_openai_available
    if not check_openai_available():
        return jsonify({"error": "OpenAI API key tidak ditemukan. Fitur ini membutuhkan OPENAI_API_KEY."}), 503

    data = request.get_json(force=True) or {}
    title = data.get("title", "").strip()
    content = data.get("content", "").strip()
    fmt = data.get("format", "berita").strip().lower()

    if not content:
        return jsonify({"error": "Konten tidak boleh kosong."}), 400
    if fmt not in FORMAT_PROMPTS:
        fmt = "berita"

    arabic_ratio = _detect_arabic_ratio(content)
    is_arabic = arabic_ratio >= 0.25   # ≥25% karakter Arab → mode Arab aktif

    try:
        client = get_openai_client()

        if is_arabic:
            # ── Mode Arab: system message khusus rekonstruksi OCR Arab ──────────
            format_instruction = _FORMAT_SYSTEM.get(fmt, _FORMAT_SYSTEM["berita"])
            user_prompt = (
                f"{format_instruction}\n\n"
                + (f"Judul/Kitab: {title}\n\n" if title else "")
                + f"=== TEKS OCR YANG PERLU DIREKONSTRUKSI ===\n{content[:7000]}\n"
                + "=== AKHIR TEKS ==="
            )
            messages = [
                {"role": "system", "content": _ARABIC_RECONSTRUCTION_SYSTEM},
                {"role": "user", "content": user_prompt},
            ]
            model = "gpt-4o"   # GPT-4o lebih akurat untuk Arab daripada gpt-4o-mini
            logger.info(f"[FORMAT-TEXT] Arab mode (ratio={arabic_ratio:.2f}), fmt={fmt}")
        else:
            # ── Mode biasa: gunakan system message AINA ──────────────────────
            prompt = FORMAT_PROMPTS[fmt](title, content[:8000])
            messages = [
                {"role": "system", "content": _AINA_SYSTEM},
                {"role": "user", "content": prompt},
            ]
            model = "gpt-4o-mini"
            logger.info(f"[FORMAT-TEXT] Non-Arab mode, fmt={fmt}")

        response = client.chat.completions.create(
            model=model,
            messages=messages,
            max_tokens=4000,
            temperature=0.1,
        )
        formatted = response.choices[0].message.content.strip()
        return jsonify({
            "status": "ok",
            "formatted_content": formatted,
            "format_used": fmt,
            "arabic_mode": is_arabic,
            "arabic_ratio": round(arabic_ratio, 2),
        })
    except Exception as e:
        logger.error(f"[FORMAT-TEXT] Error: {e}")
        return jsonify({"error": str(e)}), 500


OCR_TYPE_PROMPTS = {
    "auto": (
        "Kamu adalah OCR engine profesional multi-bahasa. Ekstrak SEMUA teks yang terlihat di gambar ini dengan akurasi tinggi.\n\n"
        "INSTRUKSI:\n"
        "- Deteksi otomatis tipe konten (poster, dokumen, kitab Arab, screenshot, dll.)\n"
        "- Ekstrak semua teks: judul, isi, tanggal, angka, nama, URL, hashtag, teks kecil\n"
        "- Pertahankan struktur & urutan (atas→bawah, kiri→kanan; untuk Arab: kanan→kiri)\n"
        "- Teks Arab: ekstrak dengan urutan RTL yang benar, pisahkan dari teks Latin\n"
        "- Pisahkan blok teks berbeda dengan satu baris kosong\n"
        "- JANGAN terjemahkan, JANGAN tambah komentar atau teks yang tidak ada di gambar\n"
        "Output: HANYA teks yang diekstrak."
    ),
    "poster": (
        "Kamu adalah OCR engine untuk poster dan flyer. Ekstrak semua teks dari gambar poster ini.\n\n"
        "INSTRUKSI:\n"
        "- Ekstrak SEMUA teks: headline besar, subtitle, body text, tanggal/waktu, lokasi, nama, nomor, URL, hashtag, teks kecil di footer\n"
        "- Urutkan dari elemen terbesar/terpenting ke terkecil\n"
        "- Pertahankan hirarki visual: judul besar → sub → detail\n"
        "- Pisahkan setiap blok elemen dengan baris kosong\n"
        "- JANGAN terjemahkan atau tambahkan teks yang tidak ada\n"
        "Output: HANYA teks dari poster, terstruktur."
    ),
    "dokumen": (
        "Kamu adalah OCR engine presisi untuk dokumen resmi. Ekstrak teks dari dokumen ini dengan akurasi maksimum.\n\n"
        "INSTRUKSI:\n"
        "- Pertahankan struktur dokumen: header, paragraf, tabel (format sebagai teks), footer\n"
        "- Ekstrak semua teks termasuk nomor dokumen, tanggal, tanda tangan tertulis, cap/stempel\n"
        "- Pertahankan indentasi dan hierarki paragraf\n"
        "- Angka dan kode penting (Nomor Surat, NIK, dll): pastikan akurasi 100%\n"
        "- Tabel: format sebagai baris dengan pemisah ` | `\n"
        "- JANGAN tambahkan interpretasi atau penjelasan\n"
        "Output: HANYA teks dokumen yang akurat."
    ),
    "kitab": (
        "Kamu adalah OCR engine spesialis teks kitab Arab dan Arab-Melayu (pegon). Ekstrak teks dari halaman kitab ini.\n\n"
        "INSTRUKSI KRITIS:\n"
        "- Teks Arab: ekstrak dengan urutan kata RTL yang BENAR, jaga harakat/tanda baca Arab\n"
        "- Teks Pegon/Arab-Melayu: ekstrak sebagaimana tertulis\n"
        "- Terjemahan/catatan Latin: ekstrak terpisah di bawah teks Arab terkait\n"
        "- Nomor halaman, nomor hadis/ayat, judul bab: sertakan semuanya\n"
        "- Pisahkan setiap paragraf/potongan teks dengan baris kosong\n"
        "- Perbaiki karakter yang terpotong/rusak berdasarkan konteks Arab\n"
        "- JANGAN terjemahkan, JANGAN tambahkan teks yang tidak ada di halaman\n"
        "Output: HANYA teks kitab yang diekstrak, Arab dan non-Arab dipertahankan."
    ),
    "screenshot": (
        "Kamu adalah OCR engine untuk screenshot digital. Ekstrak semua teks dari screenshot ini.\n\n"
        "INSTRUKSI:\n"
        "- Ekstrak semua teks: pesan, postingan, komentar, nama pengguna, timestamp, tombol, label\n"
        "- Pertahankan konteks percakapan: siapa menulis apa\n"
        "- Format: [Nama]: teks untuk percakapan/chat\n"
        "- Sertakan metadata yang terlihat: tanggal, like/share count, dll.\n"
        "- JANGAN tambahkan interpretasi atau teks yang tidak ada\n"
        "Output: HANYA teks dari screenshot, terstruktur."
    ),
}

@app.route("/api/ocr-poster", methods=["POST"])
def api_ocr_poster():
    """OCR gambar menggunakan GPT-4o Vision — multi-tipe, akurasi tinggi."""
    from ai_services import get_openai_client, check_openai_available
    import base64
    from PIL import Image as PILImage, ImageEnhance, ImageFilter
    if not check_openai_available():
        return jsonify({"error": "OpenAI API key tidak ditemukan. Fitur OCR membutuhkan OPENAI_API_KEY."}), 503
    if "image" not in request.files:
        return jsonify({"error": "Tidak ada file gambar yang dikirim."}), 400
    img_file = request.files["image"]
    if not img_file.filename:
        return jsonify({"error": "Nama file kosong."}), 400
    allowed = {"jpg", "jpeg", "png", "webp", "bmp", "gif"}
    ext = img_file.filename.rsplit(".", 1)[-1].lower() if "." in img_file.filename else ""
    if ext not in allowed:
        return jsonify({"error": f"Format file tidak didukung: .{ext}. Gunakan JPG, PNG, atau WEBP."}), 400

    ocr_type = (request.form.get("ocr_type") or "auto").strip().lower()
    if ocr_type not in OCR_TYPE_PROMPTS:
        ocr_type = "auto"

    try:
        raw = img_file.read()
        try:
            img = PILImage.open(io.BytesIO(raw)).convert("RGB")
            max_side = 1600
            if max(img.size) > max_side:
                ratio = max_side / max(img.size)
                img = img.resize((int(img.width * ratio), int(img.height * ratio)), PILImage.LANCZOS)
            if ocr_type in ("dokumen", "kitab"):
                img = ImageEnhance.Contrast(img).enhance(1.4)
                img = ImageEnhance.Sharpness(img).enhance(1.6)
            buf = io.BytesIO()
            img.save(buf, format="PNG", optimize=True)
            raw = buf.getvalue()
            mime = "image/png"
        except Exception:
            mime = "image/jpeg"
        b64 = base64.b64encode(raw).decode("utf-8")
        client = get_openai_client()
        ocr_prompt = OCR_TYPE_PROMPTS[ocr_type]
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": ocr_prompt},
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}", "detail": "high"}},
                ]
            }],
            max_tokens=3000,
            temperature=0.05,
        )
        extracted = response.choices[0].message.content.strip()
        return jsonify({"status": "ok", "text": extracted, "ocr_type": ocr_type})
    except Exception as e:
        logger.error(f"[OCR-POSTER] Error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/push-paste", methods=["POST"])
def api_push_paste():
    """Push artikel hasil Paste & Rapikan langsung ke Supabase knowledge_base."""
    from db_services import push_kb_articles
    from kb_processor import KEYWORD_TAG_MAP, DEFAULT_TAGS
    data = request.get_json(force=True) or {}
    title = data.get("title", "").strip()
    content = data.get("content", "").strip()
    if not content:
        return jsonify({"error": "Konten tidak boleh kosong."}), 400
    if not title:
        title = "Artikel dari Paste"
    # Auto-tag dari title + content
    combined = (title + " " + content).lower()
    tags = list(DEFAULT_TAGS)
    for keyword, tag in KEYWORD_TAG_MAP.items():
        if keyword in combined and tag not in tags:
            tags.append(tag)
    article = {"title": title, "content": content, "tags": tags, "summary": ""}
    try:
        result = push_kb_articles([article])
        inserted = result.get("inserted", 0)
        skipped = result.get("skipped", 0)
        if inserted > 0 or skipped > 0:
            _record_push(username=g.current_user, source="paste", count=inserted, titles=[title], skipped=skipped)
        return jsonify({
            "status": "ok",
            "inserted": inserted,
            "skipped": skipped,
            "errors": result.get("errors", []),
        })
    except Exception as e:
        logger.error(f"[PUSH-PASTE] Error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/articles/bulk-delete", methods=["POST"])
def api_articles_bulk_delete():
    """Hapus artikel terpilih berdasarkan daftar ID."""
    data = request.get_json(force=True)
    ids = data.get("ids", [])
    if not ids:
        return jsonify({"error": "Tidak ada ID yang dipilih."}), 400
    id_set = set(ids)
    articles = _load_articles()
    before = len(articles)
    articles = [a for a in articles if a.get("id") not in id_set]
    _save_articles(articles)
    deleted = before - len(articles)
    return jsonify({"status": "ok", "deleted": deleted, "remaining": len(articles)})


@app.route("/api/articles/clear-all", methods=["POST"])
def api_articles_clear_all():
    """Hapus semua hasil scraping."""
    articles = _load_articles()
    count = len(articles)
    _save_articles([])
    return jsonify({"status": "ok", "deleted": count})


@app.route("/api/reset-all", methods=["POST"])
def api_reset_all():
    """Reset semua data: artikel, KB Draft, dan progress/log."""
    global scrape_state
    # Count existing items for the response
    article_count = len(_load_articles())
    kb_count = len(_load_kb())

    # Step 1: cancel any in-flight scrape BEFORE clearing files, so the
    # background thread won't overwrite the empty files when it finishes
    with state_lock:
        scrape_state["cancelled"] = True

    # Step 2: brief pause so any _save_articles call already inside the lock
    # can finish, then our write wins cleanly
    time.sleep(0.15)

    # Step 3: now safe to clear files
    _save_articles([])
    _save_kb([])

    # Hapus file KB approved/exported jika ada
    for f in [KB_APPROVED_FILE, KB_EXPORTED_FILE]:
        try:
            if os.path.exists(f):
                os.remove(f)
        except Exception:
            pass

    # Reset progress/log and lift the cancellation flag
    with state_lock:
        scrape_state.update({
            "running": False,
            "cancelled": False,
            "phase": "idle",
            "current": 0,
            "total": 0,
            "success": 0,
            "partial": 0,
            "failed": 0,
            "duplicate": 0,
            "logs": [],
            "articles": [],
        })

    return jsonify({
        "status": "ok",
        "articles_deleted": article_count,
        "kb_deleted": kb_count,
    })


@app.route("/export/json")
def export_json():
    articles = _load_articles()
    return Response(
        json.dumps(articles, ensure_ascii=False, indent=2),
        mimetype="application/json",
        headers={"Content-Disposition": "attachment; filename=scraped_articles.json"},
    )


@app.route("/api/ingestion-history")
def api_ingestion_history():
    history_file = os.path.join(DATA_DIR, "ingestion_history.json")
    try:
        if os.path.exists(history_file):
            with open(history_file, "r", encoding="utf-8") as f:
                runs = json.load(f)
        else:
            runs = []
    except Exception:
        runs = []
    return jsonify({"runs": list(reversed(runs)), "total": len(runs)})


KB_FILE = os.path.join(DATA_DIR, "kb_articles.json")
KB_APPROVED_FILE = os.path.join(DATA_DIR, "kb_approved.json")
KB_EXPORTED_FILE = os.path.join(DATA_DIR, "kb_exported.json")
PUSH_LOG_FILE = os.path.join(DATA_DIR, "push_log.json")


def _record_push(username: str, source: str, count: int, titles: list, skipped: int = 0):
    """Catat aktivitas push ke push_log.json DAN ke Supabase (push_logs table)."""
    entry = {
        "id": f"push-{int(time.time()*1000)}",
        "timestamp": _now_iso(),
        "username": username,
        "source": source,
        "count": count,
        "skipped": skipped,
        "titles": titles[:10],
    }
    # 1. Simpan ke JSON lokal (fallback)
    try:
        log = []
        if os.path.exists(PUSH_LOG_FILE):
            with open(PUSH_LOG_FILE, "r", encoding="utf-8") as f:
                log = json.load(f)
        log.append(entry)
        with open(PUSH_LOG_FILE, "w", encoding="utf-8") as f:
            json.dump(log, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.warning(f"[PUSH-LOG] Gagal tulis JSON: {e}")
    # 2. Simpan ke Supabase (persistent, untuk production)
    try:
        from db_services import save_push_log_to_supabase
        save_push_log_to_supabase(entry)
    except Exception as e:
        logger.warning(f"[PUSH-LOG] Gagal simpan ke Supabase: {e}")

VALID_STATUSES = {"pending", "reviewed", "approved", "rejected", "exported"}
BULK_ACTION_MAP = {
    "mark_reviewed": "reviewed",
    "approve": "approved",
    "reject": "rejected",
    "export": "exported",
}


def _now_iso() -> str:
    return datetime.now().strftime("%Y-%m-%dT%H:%M:%S")


def _load_kb() -> list:
    if os.path.exists(KB_FILE):
        try:
            with open(KB_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return []


def _save_kb(data: list):
    with open(KB_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _load_file(path: str) -> list:
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return []


def _save_file(path: str, data: list):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _sync_article_to(path: str, article: dict):
    """Upsert artikel ke file JSON (approved atau exported)."""
    items = _load_file(path)
    art_id = article.get("id")
    replaced = False
    for i, item in enumerate(items):
        if item.get("id") == art_id:
            items[i] = article
            replaced = True
            break
    if not replaced:
        items.append(article)
    _save_file(path, items)


@app.route("/api/generate-summary", methods=["POST"])
def api_generate_summary():
    """Generate summary singkat untuk semua artikel scraped (tanpa AI)."""
    articles = _load_articles()
    if not articles:
        return jsonify({"error": "Belum ada artikel."}), 400

    updated = 0
    for a in articles:
        content = (a.get("content") or "").strip()
        if content and not a.get("summary"):
            a["summary"] = generate_summary(content)
            updated += 1

    _save_articles(articles)
    return jsonify({"status": "ok", "updated": updated, "total": len(articles)})


@app.route("/api/auto-tag", methods=["POST"])
def api_auto_tag():
    """Generate tags otomatis untuk semua artikel scraped (berbasis keyword)."""
    articles = _load_articles()
    if not articles:
        return jsonify({"error": "Belum ada artikel."}), 400

    for a in articles:
        title = (a.get("title") or "").strip()
        content = (a.get("content") or "").strip()
        a["tags"] = generate_tags(title, content)

    _save_articles(articles)
    return jsonify({"status": "ok", "total": len(articles)})


@app.route("/api/convert-kb", methods=["POST"])
def api_convert_kb():
    """Konversi semua artikel ke format KB draft untuk AINA.
    
    Body opsional:
      { "cutoff_days": 30 }          → hanya artikel <= 30 hari terakhir
      { "cutoff_date": "2026-01-01" } → hanya artikel >= tanggal ini
    Artikel tanpa tanggal selalu disertakan kecuali ada parameter "skip_undated": true.
    """
    data = request.get_json(force=True, silent=True) or {}
    articles = _load_articles()
    if not articles:
        return jsonify({"error": "Belum ada artikel untuk dikonversi."}), 400

    eligible = [a for a in articles if a.get("status") in ("success", "partial")]
    if not eligible:
        return jsonify({"error": "Tidak ada artikel dengan status success/partial."}), 400

    # ── Terapkan filter tanggal jika ada ──────────────────────────────────────
    cutoff: date | None = None
    if data.get("cutoff_days"):
        try:
            cutoff = date.today() - timedelta(days=int(data["cutoff_days"]))
        except (ValueError, TypeError):
            pass
    elif data.get("cutoff_date"):
        cutoff = _parse_date_param(data["cutoff_date"])

    skip_undated = bool(data.get("skip_undated", False))
    skipped = 0

    if cutoff:
        filtered = []
        for a in eligible:
            art_date = _parse_date_param(a.get("date"))
            if art_date is None:
                if skip_undated:
                    skipped += 1
                    continue
                filtered.append(a)  # tanpa tanggal → ikut sertakan
            elif art_date >= cutoff:
                filtered.append(a)
            else:
                skipped += 1
        eligible = filtered

    if not eligible:
        return jsonify({"error": "Tidak ada artikel yang memenuhi filter tanggal."}), 400

    now = _now_iso()
    existing_kb = {a["id"]: a for a in _load_kb() if a.get("id")}

    kb_articles = []
    for a in eligible:
        kb = convert_to_kb_format(a)
        prev = existing_kb.get(kb["id"], {})
        kb["approval_status"] = prev.get("approval_status", "pending")
        kb["last_updated"] = prev.get("last_updated", now)
        kb["notes"] = prev.get("notes", "")
        kb_articles.append(kb)

    _save_kb(kb_articles)
    return jsonify({"status": "ok", "count": len(kb_articles), "skipped": skipped})


@app.route("/api/kb-draft")
def api_kb_draft():
    """Ambil KB draft yang sudah dibuat."""
    kb = _load_kb()
    return jsonify({"articles": kb})


@app.route("/api/push-supabase", methods=["POST"])
def api_push_supabase():
    """Push KB articles ke AINA's Supabase knowledge_base table dengan status='pending'."""
    if not os.path.exists(KB_FILE):
        return jsonify({"error": "KB belum dikonversi. Jalankan Convert to KB Draft terlebih dahulu."}), 400
    with open(KB_FILE, "r", encoding="utf-8") as f:
        kb_articles = json.load(f)
    if not kb_articles:
        return jsonify({"error": "KB kosong."}), 400
    try:
        result = push_kb_articles(kb_articles)
        inserted = result.get("inserted", 0)
        skipped = result.get("skipped", 0)
        if inserted > 0 or skipped > 0:
            _record_push(
                username=g.current_user,
                source="review-all",
                count=inserted,
                titles=[a.get("title", "") for a in kb_articles[:10]],
                skipped=skipped,
            )
        return jsonify({
            "status": "ok",
            "inserted": inserted,
            "skipped": skipped,
            "errors": result.get("errors", []),
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/push-approved", methods=["POST"])
def api_push_approved():
    """Push hanya artikel yang sudah approved ke Supabase, lalu tandai sebagai exported."""
    if not os.path.exists(KB_APPROVED_FILE):
        return jsonify({"error": "Belum ada artikel yang diapprove."}), 400
    with open(KB_APPROVED_FILE, "r", encoding="utf-8") as f:
        approved = json.load(f)
    if not approved:
        return jsonify({"error": "Tidak ada artikel approved untuk di-push."}), 400
    try:
        result = push_kb_articles(approved)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

    inserted = result.get("inserted", 0)
    skipped = result.get("skipped", 0)
    # Mark successfully-pushed articles as exported in the main KB file
    if inserted > 0 or skipped > 0:
        _record_push(
            username=g.current_user,
            source="review-approved",
            count=inserted,
            titles=[a.get("title", "") for a in approved[:10]],
            skipped=skipped,
        )
    if inserted > 0:
        pushed_ids = {a["id"] for a in approved}
        kb = _load_kb()
        now = _now_iso()
        for a in kb:
            if a.get("id") in pushed_ids and a.get("approval_status") == "approved":
                a["approval_status"] = "exported"
                a["last_updated"] = now
                _sync_article_to(KB_EXPORTED_FILE, a)
        _save_kb(kb)
        # Clear approved file — those articles are now exported
        remaining = [a for a in approved if a.get("id") not in pushed_ids]
        with open(KB_APPROVED_FILE, "w", encoding="utf-8") as f:
            json.dump(remaining, f, ensure_ascii=False, indent=2)

    return jsonify({
        "status": "ok",
        "inserted": result["inserted"],
        "skipped": result.get("skipped", 0),
        "errors": result.get("errors", []),
    })


@app.route("/api/push-log", methods=["GET"])
def api_push_log():
    """Laporan history push ke Supabase. Hanya admin."""
    if not g.get("is_admin", False):
        return jsonify({"error": "Forbidden — hanya admin"}), 403

    # 1. Coba ambil dari Supabase dulu (persistent, tidak hilang saat redeploy)
    try:
        from db_services import fetch_push_logs_from_supabase
        supabase_log = fetch_push_logs_from_supabase(limit=200)
    except Exception:
        supabase_log = []

    if supabase_log:
        # Pastikan field skipped ada (untuk log lama yang belum punya field ini)
        for e in supabase_log:
            e.setdefault("skipped", 0)
        return jsonify({"log": supabase_log, "total": len(supabase_log), "source": "supabase"})

    # 2. Fallback ke JSON lokal (untuk dev / jika tabel push_logs belum dibuat)
    log = []
    if os.path.exists(PUSH_LOG_FILE):
        try:
            with open(PUSH_LOG_FILE, "r", encoding="utf-8") as f:
                log = json.load(f)
        except Exception:
            log = []
    for e in log:
        e.setdefault("skipped", 0)
    return jsonify({"log": list(reversed(log)), "total": len(log), "source": "local"})


@app.route("/api/push-log/clear", methods=["POST"])
def api_push_log_clear():
    """Hapus seluruh push log. Hanya admin."""
    if not g.get("is_admin", False):
        return jsonify({"error": "Forbidden — hanya admin"}), 403
    with open(PUSH_LOG_FILE, "w", encoding="utf-8") as f:
        json.dump([], f)
    return jsonify({"ok": True})


_ARAB_HEADINGS = re.compile(
    r'^\s*(بسم الله|بِسْمِ|باب|بَابٌ|بَاب|فصل|فَصْلٌ|كتاب|كِتَابٌ|مقدمة|مُقَدِّمَة|خاتمة|تمهيد|قاعدة|مسألة|الفصل|الباب|الكتاب|المقدمة|الخاتمة)',
    re.UNICODE | re.MULTILINE
)
_LATIN_HEADINGS = re.compile(
    r'^\s*(BAB|Bab|CHAPTER|Chapter|PASAL|Pasal|FASAL|Fasal|PENDAHULUAN|Pendahuluan|PENUTUP|Penutup|MUKADIMAH|Mukadimah|PROLOG|EPILOG)\b',
    re.MULTILINE
)

@app.route("/api/pdf/inspect", methods=["POST"])
def api_pdf_inspect():
    """
    Inspeksi PDF: baca metadata, judul, dan daftar halaman dengan deteksi bab/fasal/heading.
    Kembalikan ringkasan detail sebelum ekstraksi dimulai.
    """
    if "file" not in request.files:
        return jsonify({"error": "Tidak ada file yang dikirim."}), 400
    pdf_file = request.files["file"]
    if not pdf_file.filename or not pdf_file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "File harus berformat PDF."}), 400

    try:
        pdf_bytes = pdf_file.read()

        # ── Buka dengan PyMuPDF untuk metadata ──
        fitz_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_pages = len(fitz_doc)
        meta = fitz_doc.metadata or {}
        pdf_title = (meta.get("title") or "").strip()
        pdf_author = (meta.get("author") or "").strip()
        pdf_subject = (meta.get("subject") or "").strip()
        fitz_doc.close()

        # ── Buka dengan pdfplumber untuk analisis per halaman ──
        pages_info = []
        chapters = []   # list of {page_num, heading, heading_type}
        text_count = 0
        scan_count = 0

        MAX_INSPECT_PAGES = 800  # cap agar tidak OOM

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            inspect_pages = min(total_pages, MAX_INSPECT_PAGES)
            for i in range(inspect_pages):
                page = pdf.pages[i]
                raw = (page.extract_text() or "").strip()
                page_num = i + 1
                is_scan = not raw or len(raw) < 15

                if is_scan:
                    scan_count += 1
                    pages_info.append({
                        "page": page_num,
                        "type": "scan",
                        "heading": None,
                        "preview": "[Halaman gambar/scan — perlu OCR]",
                        "words": 0,
                    })
                    continue

                text_count += 1
                word_count = len(raw.split())
                first_line = raw.split('\n')[0][:120].strip()

                # Deteksi heading Arabic atau Latin
                heading = None
                heading_type = None
                arab_match = _ARAB_HEADINGS.search(raw[:300])
                latin_match = _LATIN_HEADINGS.search(raw[:300])
                if arab_match:
                    # Ambil baris pertama yang mengandung heading Arab
                    for ln in raw.split('\n')[:5]:
                        if _ARAB_HEADINGS.search(ln):
                            heading = ln.strip()[:120]
                            break
                    heading_type = "arab"
                elif latin_match:
                    for ln in raw.split('\n')[:5]:
                        if _LATIN_HEADINGS.search(ln):
                            heading = ln.strip()[:120]
                            break
                    heading_type = "latin"

                if heading:
                    chapters.append({"page": page_num, "heading": heading, "type": heading_type})

                pages_info.append({
                    "page": page_num,
                    "type": "text",
                    "heading": heading,
                    "heading_type": heading_type,
                    "preview": first_line,
                    "words": word_count,
                })

        # Halaman yang tidak di-inspect (> MAX_INSPECT_PAGES) dianggap scan
        if total_pages > MAX_INSPECT_PAGES:
            for i in range(MAX_INSPECT_PAGES, total_pages):
                pages_info.append({
                    "page": i + 1,
                    "type": "unknown",
                    "heading": None,
                    "preview": "[Tidak di-inspeksi — di luar batas preview]",
                    "words": 0,
                })

        # Nama file sebagai fallback judul
        base_name = os.path.splitext(pdf_file.filename)[0].replace("-", " ").replace("_", " ").strip()

        return jsonify({
            "ok": True,
            "filename": pdf_file.filename,
            "title": pdf_title or base_name,
            "author": pdf_author,
            "subject": pdf_subject,
            "total_pages": total_pages,
            "text_pages": text_count,
            "scan_pages": scan_count,
            "chapters": chapters,
            "pages": pages_info,
        })

    except Exception as e:
        logger.error(f"[PDF INSPECT] Error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/pdf/learn", methods=["POST"])
def api_pdf_learn():
    """
    Analisis isi PDF dengan AI: ringkasan keseluruhan + deskripsi tiap bab/fasal.
    Menggunakan GPT-4o untuk membaca sampel teks dan menjelaskan pembahasannya.
    """
    if "file" not in request.files:
        return jsonify({"error": "Tidak ada file yang dikirim."}), 400
    pdf_file = request.files["file"]
    if not pdf_file.filename or not pdf_file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "File harus berformat PDF."}), 400
    if not check_openai_available():
        return jsonify({"error": "OPENAI_API_KEY tidak tersedia — fitur analisis AI tidak aktif."}), 503

    try:
        pdf_bytes = pdf_file.read()
        base_name = os.path.splitext(pdf_file.filename)[0].replace("-", " ").replace("_", " ").strip()

        # ── Ekstrak teks per halaman (hingga 300 hal, hemat memori) ──
        pages_text_raw = []   # list of (page_num, text)
        chapters_detected = []

        fitz_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_pages = len(fitz_doc)
        meta = fitz_doc.metadata or {}
        pdf_title = (meta.get("title") or "").strip() or base_name
        fitz_doc.close()

        MAX_PAGES = min(total_pages, 300)

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for i in range(MAX_PAGES):
                page = pdf.pages[i]
                raw = (page.extract_text() or "").strip()
                if raw and len(raw) > 15:
                    pages_text_raw.append((i + 1, raw))
                    # Deteksi chapter headings
                    arab_m = _ARAB_HEADINGS.search(raw[:400])
                    latin_m = _LATIN_HEADINGS.search(raw[:400])
                    heading = None
                    if arab_m:
                        for ln in raw.split('\n')[:6]:
                            if _ARAB_HEADINGS.search(ln):
                                heading = ln.strip()[:150]
                                break
                    elif latin_m:
                        for ln in raw.split('\n')[:6]:
                            if _LATIN_HEADINGS.search(ln):
                                heading = ln.strip()[:150]
                                break
                    if heading:
                        chapters_detected.append({"page": i + 1, "heading": heading, "text_sample": raw[:600]})

        text_pages = len(pages_text_raw)
        scan_pages = MAX_PAGES - text_pages

        if text_pages == 0:
            return jsonify({
                "ok": True,
                "title": pdf_title,
                "total_pages": total_pages,
                "text_pages": 0,
                "scan_pages": scan_pages,
                "ai_available": False,
                "overview": "PDF ini tampaknya berisi halaman scan/gambar semua — tidak ada teks digital yang bisa dianalisis. Aktifkan OCR untuk mengekstrak teksnya terlebih dahulu.",
                "chapters": [],
                "topics": [],
            })

        # ── Siapkan sampel teks untuk AI ──
        # Ambil: 5 hal pertama + 5 hal terakhir + tiap chapter + max 15 hal acak
        sample_pages = set()
        for i in range(min(5, len(pages_text_raw))):
            sample_pages.add(i)
        for i in range(max(0, len(pages_text_raw) - 5), len(pages_text_raw)):
            sample_pages.add(i)
        # Tiap chapter
        ch_indices = {ch["page"] - 1 for ch in chapters_detected if ch["page"] - 1 < len(pages_text_raw)}
        sample_pages.update(ch_indices)
        # Tambah sampel merata
        step = max(1, len(pages_text_raw) // 15)
        for i in range(0, len(pages_text_raw), step):
            sample_pages.add(i)

        sampled = sorted(sample_pages)[:35]
        sample_text_parts = []
        for idx in sampled:
            pnum, ptxt = pages_text_raw[idx]
            snippet = ptxt[:500].strip()
            sample_text_parts.append(f"[Halaman {pnum}]\n{snippet}")
        sample_text = "\n\n---\n\n".join(sample_text_parts)

        # ── Siapkan daftar bab untuk AI ──
        chapters_for_prompt = "\n".join(
            f"- Halaman {ch['page']}: {ch['heading']}"
            for ch in chapters_detected[:30]
        ) or "(tidak terdeteksi otomatis)"

        # ── Bangun prompt AI ──
        prompt = f"""Kamu adalah asisten analisis kitab/buku berbahasa Arab dan Indonesia.

Berikut adalah metadata dan sampel teks dari PDF berjudul "{pdf_title}":
- Total halaman: {total_pages} (teks digital: {text_pages}, scan: {scan_pages})
- Bab/fasal terdeteksi otomatis:
{chapters_for_prompt}

=== SAMPEL TEKS PDF ===
{sample_text[:6000]}
=== AKHIR SAMPEL ===

Berdasarkan informasi di atas, buatlah analisis dalam format JSON berikut (jawab HANYA JSON, tanpa markdown):
{{
  "judul": "judul/nama kitab yang sebenarnya (dari teks)",
  "penulis": "nama penulis jika ditemukan, kosong jika tidak",
  "bahasa": "Arab / Indonesia / Campuran",
  "bidang": "bidang ilmu (misal: Fiqh, Aqidah, Tafsir, Hadits, dsb)",
  "overview": "ringkasan 2-3 kalimat tentang isi dan tujuan kitab ini",
  "bab": [
    {{
      "nomor": 1,
      "judul": "judul bab/fasal (teks asli jika Arab)",
      "halaman": 1,
      "pembahasan": "1-2 kalimat tentang apa yang dibahas di bab ini"
    }}
  ],
  "topik_utama": ["topik1", "topik2", "topik3"]
}}

Jika bab tidak terdeteksi otomatis, identifikasi dari konten teks. Maksimal 20 bab."""

        client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY", ""))
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=2000,
            temperature=0.3,
        )
        raw_json = response.choices[0].message.content.strip()

        # Bersihkan jika ada markdown fence
        if raw_json.startswith("```"):
            raw_json = re.sub(r"^```[a-z]*\n?", "", raw_json)
            raw_json = re.sub(r"\n?```$", "", raw_json)

        ai_result = json.loads(raw_json)

        return jsonify({
            "ok": True,
            "title": ai_result.get("judul") or pdf_title,
            "author": ai_result.get("penulis", ""),
            "language": ai_result.get("bahasa", ""),
            "field": ai_result.get("bidang", ""),
            "total_pages": total_pages,
            "text_pages": text_pages,
            "scan_pages": scan_pages,
            "ai_available": True,
            "overview": ai_result.get("overview", ""),
            "chapters": ai_result.get("bab", []),
            "topics": ai_result.get("topik_utama", []),
        })

    except json.JSONDecodeError as je:
        logger.error(f"[PDF LEARN] JSON parse error: {je}")
        return jsonify({"error": "AI mengembalikan format tidak valid. Coba lagi."}), 500
    except Exception as e:
        logger.error(f"[PDF LEARN] Error: {e}")
        return jsonify({"error": str(e)}), 500


def _safe_pdf_slug(filename: str, suffix: str = "") -> str:
    """Buat slug aman dari nama file PDF (bisa non-ASCII/Arab)."""
    base = os.path.splitext(filename)[0]
    # Try ASCII slug first
    slug = generate_slug(base)
    if not slug:
        # Fallback: hex of filename + suffix
        slug = re.sub(r"[^a-z0-9]", "-", base.lower())[:30].strip("-") or "kitab"
    if suffix:
        slug = f"{slug}-{suffix}"
    return slug[:80]


# ── PDF Job store (in-memory, background processing) ──────────────────────────
import threading as _threading
import uuid as _uuid

_pdf_jobs: dict = {}           # job_id → job dict
_pdf_jobs_lock = _threading.Lock()


def _pdf_job_worker(job_id: str, files_data: list, category: str, chunk_size: int,
                    use_ocr: bool, max_ocr_pages: int, page_start_g: int, page_end_g: int):
    """Background thread: process PDF files and store results in _pdf_jobs."""
    OCR_BATCH = 4
    ocr_available = check_openai_available()
    all_results = []
    total_files = len(files_data)

    def _set_progress(msg: str, done_files: int = 0):
        with _pdf_jobs_lock:
            _pdf_jobs[job_id]["progress"] = msg
            _pdf_jobs[job_id]["done_files"] = done_files

    # Jika OCR diminta tapi API key tidak tersedia, langsung beri tahu
    if use_ocr and not ocr_available:
        with _pdf_jobs_lock:
            _pdf_jobs[job_id].update({
                "status": "done",
                "progress": "OCR tidak bisa dijalankan — OPENAI_API_KEY belum diset.",
                "done_files": 0,
                "results": {
                    "status": "ok",
                    "processed": 0,
                    "total": total_files,
                    "total_chunks": 0,
                    "results": [{"filename": f, "status": "error",
                                 "error": "OPENAI_API_KEY tidak ditemukan. Set secret OPENAI_API_KEY di environment untuk menggunakan OCR."}
                                for f, _ in files_data],
                },
            })
        logger.warning(f"[PDF JOB] {job_id} dibatalkan — OPENAI_API_KEY tidak ada")
        return

    _set_progress(f"Memulai pemrosesan {total_files} file...")

    try:
        for fi, (fname, pdf_bytes) in enumerate(files_data):
            _set_progress(f"[{fi+1}/{total_files}] Membaca {fname}...", fi)
            try:
                pages_text = []
                scan_page_indices = []
                text_pages = 0

                fitz_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                total_pages = len(fitz_doc)
                p_start_idx = max(0, page_start_g - 1)
                p_end_idx = min(total_pages - 1, page_end_g - 1) if page_end_g > 0 else total_pages - 1

                with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                    for i, page in enumerate(pdf.pages):
                        if i < p_start_idx or i > p_end_idx:
                            continue
                        raw_text = (page.extract_text() or "").strip()
                        if raw_text and len(raw_text) > 20:
                            pages_text.append((i + 1, raw_text, False))
                            text_pages += 1
                        else:
                            pages_text.append((i + 1, "", True))
                            scan_page_indices.append(i)

                scan_pages = len(scan_page_indices)

                if use_ocr and ocr_available and scan_page_indices:
                    ocr_indices = scan_page_indices[:max_ocr_pages]
                    ocr_skipped = len(scan_page_indices) - len(ocr_indices)
                    total_batches = (len(ocr_indices) + OCR_BATCH - 1) // OCR_BATCH

                    if ocr_skipped > 0:
                        logger.info(f"[PDF OCR] Cap: {ocr_skipped} hal. skip (limit {max_ocr_pages})")

                    mat = fitz.Matrix(1.5, 1.5)
                    scan_imgs = {}
                    for i in ocr_indices:
                        pix = fitz_doc[i].get_pixmap(matrix=mat)
                        scan_imgs[i] = pix.tobytes("png")

                    idx_list = list(scan_imgs.keys())
                    for bn, batch_start_idx in enumerate(range(0, len(idx_list), OCR_BATCH)):
                        batch_end_idx = min(batch_start_idx + OCR_BATCH - 1, len(idx_list) - 1)
                        _set_progress(
                            f"[{fi+1}/{total_files}] OCR {fname}: batch {bn+1}/{total_batches} "
                            f"(hal. {idx_list[batch_start_idx]+1}–{idx_list[batch_end_idx]+1})",
                            fi,
                        )
                        batch_idx = idx_list[batch_start_idx:batch_start_idx + OCR_BATCH]
                        batch_imgs = [scan_imgs[j] for j in batch_idx]
                        try:
                            ocr_texts = ocr_arabic_pages_batch(batch_imgs)
                            for j, ocr_text in zip(batch_idx, ocr_texts):
                                for k, (pnum, _, is_scan) in enumerate(pages_text):
                                    if pnum == j + 1 and is_scan:
                                        pages_text[k] = (pnum, ocr_text or "", True)
                                        break
                        except Exception as ocr_err:
                            logger.warning(f"[PDF OCR] Batch error: {ocr_err}")

                fitz_doc.close()

                _set_progress(f"[{fi+1}/{total_files}] Menyimpan chunk KB untuk {fname}...", fi)
                base_title = os.path.splitext(fname)[0].replace("-", " ").replace("_", " ").strip()
                extra_tags = [category] if category else []
                chunks_created = 0
                ts = int(time.time())
                now = _now_iso()
                kb = _load_kb()

                for _ci, start in enumerate(range(0, len(pages_text), chunk_size)):
                    chunk_pages = pages_text[start:start + chunk_size]
                    pg_start = chunk_pages[0][0]
                    pg_end = chunk_pages[-1][0]
                    chunk_text = "\n\n".join(t for _, t, _ in chunk_pages if t).strip()
                    if not chunk_text:
                        continue

                    if total_pages > chunk_size:
                        chunk_title = f"{base_title} — Hal. {pg_start}–{pg_end}"
                        slug_suffix = f"hal{pg_start}-{pg_end}-{ts}"
                    else:
                        chunk_title = base_title
                        slug_suffix = str(ts)

                    slug = _safe_pdf_slug(fname, slug_suffix)
                    tags = list(set(extra_tags + generate_tags(chunk_title, chunk_text[:600])))
                    summary = generate_summary(chunk_title, chunk_text)
                    article_id = f"pdf-{slug}"

                    kb_article = {
                        "id": article_id,
                        "title": chunk_title,
                        "slug": slug,
                        "source_url": f"pdf://{fname}",
                        "published_date": now[:10],
                        "content": chunk_text,
                        "summary": summary,
                        "tags": tags,
                        "scrape_status": "success",
                        "approval_status": "pending",
                        "last_updated": now,
                        "notes": f"Sumber: {fname}, Hal. {pg_start}-{pg_end}",
                        "source_type": "pdf",
                    }

                    idx = next((j for j, a in enumerate(kb) if a.get("id") == article_id), None)
                    if idx is not None:
                        kb[idx] = kb_article
                    else:
                        kb.append(kb_article)
                    chunks_created += 1

                _save_kb(kb)

                ocr_pages_done = len(scan_page_indices[:max_ocr_pages]) if (use_ocr and ocr_available) else 0
                all_results.append({
                    "filename": fname,
                    "status": "ok" if chunks_created > 0 else "error",
                    "error": "Tidak ada teks yang bisa diekstrak." if chunks_created == 0 else None,
                    "title": base_title,
                    "total_pages": total_pages,
                    "text_pages": text_pages,
                    "scan_pages": scan_pages,
                    "ocr_pages_done": ocr_pages_done,
                    "chunks": chunks_created,
                })

            except Exception as e:
                logger.error(f"[PDF JOB] Error on {fname}: {e}", exc_info=True)
                all_results.append({"filename": fname, "status": "error", "error": str(e)[:300]})

    except Exception as outer_e:
        logger.error(f"[PDF JOB] Outer error: {outer_e}", exc_info=True)
        all_results.append({"filename": "—", "status": "error", "error": f"Error tak terduga: {str(outer_e)[:200]}"})

    finally:
        ok_count = sum(1 for r in all_results if r.get("status") == "ok")
        total_chunks = sum(r.get("chunks", 0) for r in all_results)

        with _pdf_jobs_lock:
            _pdf_jobs[job_id].update({
                "status": "done",
                "progress": f"Selesai — {ok_count}/{total_files} file berhasil, {total_chunks} chunk disimpan.",
                "done_files": total_files,
                "results": {
                    "status": "ok",
                    "processed": ok_count,
                    "total": total_files,
                    "total_chunks": total_chunks,
                    "results": all_results,
                },
            })
        logger.info(f"[PDF JOB] {job_id} selesai — {ok_count}/{total_files} ok, {total_chunks} chunks")


@app.route("/api/pdf/upload", methods=["POST"])
def api_pdf_upload():
    """
    Upload satu atau lebih PDF — langsung kembalikan job_id, proses di background.
    Polling status via GET /api/pdf/job/<job_id>
    """
    files = request.files.getlist("files")
    if not files or all(f.filename == "" for f in files):
        return jsonify({"error": "Tidak ada file PDF yang diupload."}), 400

    category = request.form.get("category", "").strip()
    chunk_size = max(5, min(100, int(request.form.get("chunk_size", "20"))))
    use_ocr = request.form.get("use_ocr", "false").lower() == "true"
    max_ocr_pages = max(10, min(500, int(request.form.get("max_ocr_pages", "150"))))
    page_start_g = max(1, int(request.form.get("page_start", "1") or "1"))
    page_end_raw = request.form.get("page_end", "0") or "0"
    page_end_g = int(page_end_raw) if page_end_raw.isdigit() and int(page_end_raw) > 0 else 0

    # Baca bytes semua file sebelum keluar dari request context
    files_data = []
    for f in files:
        fname = f.filename or ""
        if not fname.lower().endswith(".pdf"):
            continue
        files_data.append((fname, f.read()))

    if not files_data:
        return jsonify({"error": "Tidak ada file PDF yang valid."}), 400

    job_id = str(_uuid.uuid4())
    with _pdf_jobs_lock:
        _pdf_jobs[job_id] = {
            "status": "processing",
            "progress": "Menginisialisasi...",
            "done_files": 0,
            "total_files": len(files_data),
            "results": None,
            "created_at": time.time(),
        }

    t = _threading.Thread(
        target=_pdf_job_worker,
        args=(job_id, files_data, category, chunk_size, use_ocr, max_ocr_pages, page_start_g, page_end_g),
        daemon=True,
    )
    t.start()
    logger.info(f"[PDF JOB] Mulai job {job_id} — {len(files_data)} file, OCR={use_ocr}")

    return jsonify({"job_id": job_id, "status": "processing", "total_files": len(files_data)})


@app.route("/api/pdf/job/<job_id>", methods=["GET"])
def api_pdf_job_status(job_id: str):
    """Poll status job upload PDF."""
    with _pdf_jobs_lock:
        job = _pdf_jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job tidak ditemukan."}), 404
    # Cleanup job lama (lebih dari 1 jam)
    if time.time() - job.get("created_at", 0) > 3600:
        with _pdf_jobs_lock:
            _pdf_jobs.pop(job_id, None)
        return jsonify({"error": "Job expired."}), 404
    return jsonify({
        "job_id": job_id,
        "status": job["status"],
        "progress": job.get("progress", ""),
        "done_files": job.get("done_files", 0),
        "total_files": job.get("total_files", 1),
        "results": job.get("results"),   # None saat masih proses
    })


@app.route("/api/pdf/rapikan-file", methods=["POST"])
def api_pdf_rapikan_file():
    """Perbaiki semua chunk KB dari satu file PDF dengan AI (bersihkan teks OCR, format ulang)."""
    from ai_services import get_openai_client, check_openai_available
    if not check_openai_available():
        return jsonify({"error": "OpenAI API key tidak ditemukan."}), 503
    data = request.get_json(force=True) or {}
    filename = data.get("filename", "").strip()
    if not filename:
        return jsonify({"error": "Nama file diperlukan."}), 400
    kb = _load_kb()
    target = [a for a in kb if a.get("source_url", "").startswith(f"pdf://{filename}")]
    if not target:
        return jsonify({"error": f"Tidak ada KB chunk ditemukan untuk file: {filename}"}), 404
    client = get_openai_client()
    updated = 0
    for article in target:
        raw_content = (article.get("content") or "").strip()
        if not raw_content:
            continue
        try:
            prompt = (
                "Kamu adalah editor teks Arab/Indonesia profesional. Perbaiki dan rapikan teks berikut yang merupakan hasil ekstraksi dari PDF kitab.\n\n"
                "INSTRUKSI:\n"
                "- Perbaiki kesalahan OCR: karakter aneh, spasi yang salah, baris terputus\n"
                "- Pertahankan semua konten asli — JANGAN hapus atau tambah informasi\n"
                "- Rapikan struktur paragraf dan spasi\n"
                "- Pertahankan bahasa asli (Arab, Indonesia, atau campuran)\n"
                "- Jika ada teks Arab: pastikan urutan kata Arab tetap RTL\n"
                "- Output: HANYA teks yang sudah diperbaiki, tanpa komentar apapun\n\n"
                f"Teks:\n{raw_content[:4000]}"
            )
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2000,
                temperature=0.05,
            )
            fixed = resp.choices[0].message.content.strip()
            if fixed:
                article["content"] = fixed
                article["last_updated"] = _now_iso()
                updated += 1
        except Exception as e:
            logger.warning(f"[PDF-RAPIKAN] Gagal chunk {article.get('id')}: {e}")
    if updated > 0:
        _save_kb(kb)
    return jsonify({"status": "ok", "updated": updated, "total": len(target)})


@app.route("/api/ai-summary", methods=["POST"])
def api_ai_summary():
    """Generate AI summary untuk satu artikel."""
    data = request.get_json(force=True)
    article_id = (data.get("id") or "").strip()
    if not article_id:
        return jsonify({"error": "ID artikel diperlukan."}), 400

    articles = _load_articles()
    article = next((a for a in articles if a["id"] == article_id), None)
    if not article:
        return jsonify({"error": "Artikel tidak ditemukan."}), 404

    try:
        summary = generate_ai_summary(
            title=article.get("title") or "",
            content=article.get("content") or "",
        )
        return jsonify({"status": "ok", "summary": summary})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ai-summary-all", methods=["POST"])
def api_ai_summary_all():
    """Generate AI summary (GPT-4o-mini) untuk semua artikel KB dan update KB file."""
    kb_articles = _load_kb()
    if not kb_articles:
        return jsonify({"error": "KB belum dikonversi. Jalankan Convert to KB Draft terlebih dahulu."}), 400

    errors = []
    for art in kb_articles:
        try:
            art["ai_summary"] = generate_ai_summary(
                title=art.get("title") or "",
                content=art.get("content") or "",
            )
            # Update summary field juga jika AI berhasil
            art["summary"] = art["ai_summary"]
        except Exception as e:
            errors.append({"slug": art.get("slug"), "error": str(e)})

    _save_kb(kb_articles)
    return jsonify({"status": "ok", "count": len(kb_articles), "errors": errors})


@app.route("/api/db-articles")
def api_db_articles():
    """Ambil semua artikel dari Supabase."""
    try:
        articles = fetch_kb_articles_from_db()
        return jsonify(articles)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/export/kb")
def export_kb():
    if not os.path.exists(KB_FILE):
        return jsonify({"error": "KB belum dikonversi."}), 404
    with open(KB_FILE, "r", encoding="utf-8") as f:
        data = f.read()
    return Response(
        data,
        mimetype="application/json",
        headers={"Content-Disposition": "attachment; filename=kb_articles.json"},
    )


@app.route("/export/csv")
def export_csv():
    articles = _load_articles()
    si = io.StringIO()
    fields = ["id", "title", "date", "url", "content", "status", "error_reason", "mode"]
    writer = csv.DictWriter(si, fieldnames=fields, extrasaction="ignore")
    writer.writeheader()
    writer.writerows(articles)
    return Response(
        si.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=scraped_articles.csv"},
    )


# ─── KB Review Workflow ───────────────────────────────────────────────────────

@app.route("/kb-drafts")
def kb_drafts():
    """Ambil semua KB draft, dengan optional filter ?status=pending|reviewed|approved|rejected|exported."""
    status_filter = request.args.get("status", "").strip().lower()
    kb = _load_kb()
    if status_filter and status_filter != "all":
        kb = [a for a in kb if a.get("approval_status") == status_filter]
    return jsonify(kb)


@app.route("/kb/update-status", methods=["POST"])
def kb_update_status():
    """Update status, notes, content, dan/atau summary satu artikel KB."""
    data = request.get_json(force=True)
    article_id = (data.get("id") or "").strip()
    new_status = (data.get("status") or "").strip()
    notes = data.get("notes")
    new_content = data.get("content")
    new_summary = data.get("summary")

    if new_status not in VALID_STATUSES:
        return jsonify({"error": f"Status tidak valid: {new_status}. Pilih: {', '.join(VALID_STATUSES)}"}), 400

    kb = _load_kb()
    article = next((a for a in kb if a.get("id") == article_id), None)
    if not article:
        return jsonify({"error": "Artikel tidak ditemukan"}), 404

    article["approval_status"] = new_status
    article["last_updated"] = _now_iso()
    if notes is not None:
        article["notes"] = str(notes)
    if new_content is not None:
        article["content"] = str(new_content)
    if new_summary is not None:
        article["summary"] = str(new_summary)

    _save_kb(kb)

    # Sync ke file approved/exported
    if new_status == "approved":
        _sync_article_to(KB_APPROVED_FILE, article)
    elif new_status == "exported":
        _sync_article_to(KB_EXPORTED_FILE, article)

    return jsonify({"status": "ok", "article": article})


@app.route("/kb/bulk-action", methods=["POST"])
def kb_bulk_action():
    """Ubah status banyak artikel sekaligus."""
    data = request.get_json(force=True)
    ids = data.get("ids", [])
    action = (data.get("action") or "").strip()

    if action not in BULK_ACTION_MAP:
        return jsonify({"error": f"Action tidak valid: {action}. Pilih: {', '.join(BULK_ACTION_MAP)}"}), 400
    if not ids:
        return jsonify({"error": "Tidak ada ID yang dipilih"}), 400

    new_status = BULK_ACTION_MAP[action]
    id_set = set(ids)
    kb = _load_kb()
    now = _now_iso()
    updated = 0

    for a in kb:
        if a.get("id") in id_set:
            a["approval_status"] = new_status
            a["last_updated"] = now
            updated += 1
            if new_status == "approved":
                _sync_article_to(KB_APPROVED_FILE, a)
            elif new_status == "exported":
                _sync_article_to(KB_EXPORTED_FILE, a)

    _save_kb(kb)
    return jsonify({"status": "ok", "updated": updated, "new_status": new_status})


@app.route("/kb/stats")
def kb_stats():
    """Hitung jumlah artikel per status."""
    kb = _load_kb()
    counts: dict[str, int] = {s: 0 for s in VALID_STATUSES}
    counts["total"] = len(kb)
    for a in kb:
        s = a.get("approval_status", "pending")
        if s in counts:
            counts[s] += 1
    return jsonify(counts)


@app.route("/kb/delete", methods=["POST"])
def kb_delete():
    """Hapus satu artikel KB berdasarkan id."""
    data = request.get_json(force=True)
    article_id = (data.get("id") or "").strip()
    if not article_id:
        return jsonify({"error": "id wajib diisi"}), 400

    kb = _load_kb()
    new_kb = [a for a in kb if a.get("id") != article_id]
    if len(new_kb) == len(kb):
        return jsonify({"error": "Artikel tidak ditemukan"}), 404

    _save_kb(new_kb)

    # Hapus dari file approved/exported juga
    for fpath in [KB_APPROVED_FILE, KB_EXPORTED_FILE]:
        try:
            items = _load_file(fpath)
            filtered = [a for a in items if a.get("id") != article_id]
            if len(filtered) != len(items):
                with open(fpath, "w", encoding="utf-8") as f:
                    json.dump(filtered, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    return jsonify({"status": "ok", "deleted": article_id})


@app.route("/kb/bulk-delete", methods=["POST"])
def kb_bulk_delete():
    """Hapus banyak artikel KB sekaligus berdasarkan array ids."""
    data = request.get_json(force=True)
    ids = data.get("ids", [])
    if not ids:
        return jsonify({"error": "Tidak ada ID yang dipilih"}), 400

    id_set = set(ids)
    kb = _load_kb()
    new_kb = [a for a in kb if a.get("id") not in id_set]
    deleted = len(kb) - len(new_kb)
    _save_kb(new_kb)

    for fpath in [KB_APPROVED_FILE, KB_EXPORTED_FILE]:
        try:
            items = _load_file(fpath)
            filtered = [a for a in items if a.get("id") not in id_set]
            if len(filtered) != len(items):
                with open(fpath, "w", encoding="utf-8") as f:
                    json.dump(filtered, f, ensure_ascii=False, indent=2)
        except Exception:
            pass

    return jsonify({"status": "ok", "deleted": deleted})


@app.route("/kb/reset", methods=["POST"])
def kb_reset():
    """Hapus semua KB draft (reset total)."""
    _save_kb([])
    for fpath in [KB_APPROVED_FILE, KB_EXPORTED_FILE]:
        try:
            with open(fpath, "w", encoding="utf-8") as f:
                json.dump([], f)
        except Exception:
            pass
    return jsonify({"status": "ok"})


@app.route("/export/kb-approved")
def export_kb_approved():
    items = _load_file(KB_APPROVED_FILE)
    return Response(
        json.dumps(items, ensure_ascii=False, indent=2),
        mimetype="application/json",
        headers={"Content-Disposition": "attachment; filename=kb_approved.json"},
    )


@app.route("/export/kb-exported")
def export_kb_exported():
    items = _load_file(KB_EXPORTED_FILE)
    return Response(
        json.dumps(items, ensure_ascii=False, indent=2),
        mimetype="application/json",
        headers={"Content-Disposition": "attachment; filename=kb_exported.json"},
    )


@app.route("/export/kb-markdown")
def export_kb_markdown():
    """Export semua KB draft sebagai file .md yang rapi dan mudah dibaca."""
    kb = _load_kb()
    if not kb:
        return jsonify({"error": "Belum ada KB Draft."}), 404

    STATUS_ICON = {
        "pending": "⏳",
        "reviewed": "🔍",
        "approved": "✅",
        "rejected": "❌",
        "exported": "📦",
    }

    today = datetime.now().strftime("%d %B %Y, %H:%M")
    lines: list[str] = []

    # ── Header dokumen ──────────────────────────────────────────────────────
    lines.append("# Knowledge Base — AINA")
    lines.append("")
    lines.append(f"> **Diekspor pada:** {today}  ")
    lines.append(f"> **Total artikel:** {len(kb)}")
    lines.append("")
    lines.append("---")
    lines.append("")

    for i, article in enumerate(kb, 1):
        title = article.get("title") or "Tanpa Judul"
        slug = article.get("slug") or "-"
        source_url = article.get("source_url") or ""
        pub_date = article.get("published_date") or "-"
        summary = (article.get("summary") or "").strip()
        content = (article.get("content") or "").strip()
        tags = article.get("tags") or []
        status = article.get("approval_status") or "pending"
        status_display = f"{STATUS_ICON.get(status, '')} {status}"

        # ── Judul artikel ──────────────────────────────────────────────────
        lines.append(f"## {i}. {title}")
        lines.append("")

        # ── Metadata ───────────────────────────────────────────────────────
        lines.append("| | |")
        lines.append("|:---|:---|")
        lines.append(f"| **Tanggal terbit** | {pub_date} |")
        if source_url:
            lines.append(f"| **Sumber** | [{source_url}]({source_url}) |")
        lines.append(f"| **Slug** | `{slug}` |")
        if tags:
            tag_str = " · ".join(f"`{t}`" for t in tags)
            lines.append(f"| **Tags** | {tag_str} |")
        lines.append(f"| **Status** | {status_display} |")
        lines.append("")

        # ── Ringkasan ──────────────────────────────────────────────────────
        if summary:
            lines.append("### Ringkasan")
            lines.append("")
            lines.append(summary)
            lines.append("")

        # ── Konten Lengkap ─────────────────────────────────────────────────
        if content:
            lines.append("### Konten Lengkap")
            lines.append("")
            # Bagi konten menjadi paragraf berdasarkan newline
            for para in content.split("\n"):
                para = para.strip()
                if para:
                    lines.append(para)
                    lines.append("")

        lines.append("---")
        lines.append("")

    md_output = "\n".join(lines)
    filename = f"kb_aina_{datetime.now().strftime('%Y%m%d_%H%M')}.md"
    return Response(
        md_output,
        mimetype="text/markdown; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# ─── Scheduler API ───────────────────────────────────────────────────────────

@app.route("/api/scheduler/settings", methods=["GET"])
def scheduler_get_settings():
    cfg = _load_scheduler_settings()
    cfg["next_run_at"] = _next_run_iso()
    return jsonify(cfg)


@app.route("/api/scheduler/settings", methods=["POST"])
def scheduler_post_settings():
    data = request.get_json(force=True)
    cfg = _load_scheduler_settings()

    allowed = {"enabled", "interval", "day_of_week", "time_of_day",
               "url", "scrape_mode", "incremental", "rss_sources"}
    for k in allowed:
        if k in data:
            cfg[k] = data[k]

    # Validasi interval
    if cfg["interval"] not in ("manual", "daily", "weekly"):
        cfg["interval"] = "manual"
    if cfg["interval"] == "manual":
        cfg["enabled"] = False

    # Validasi rss_sources
    if "rss_sources" in cfg:
        clean = []
        for src in cfg["rss_sources"]:
            u = (src.get("url") or "").strip()
            if u:
                clean.append({
                    "url": u,
                    "label": (src.get("label") or "").strip() or u,
                    "max_items": max(1, min(100, int(src.get("max_items") or 10))),
                })
        cfg["rss_sources"] = clean

    _save_scheduler_settings(cfg)
    _apply_scheduler(cfg)

    cfg["next_run_at"] = _next_run_iso()
    return jsonify({"status": "ok", "settings": cfg})


@app.route("/api/scheduler/rss-sources", methods=["GET"])
def scheduler_rss_sources_get():
    cfg = _load_scheduler_settings()
    return jsonify({"rss_sources": cfg.get("rss_sources", [])})


@app.route("/api/scheduler/rss-sources", methods=["POST"])
def scheduler_rss_sources_post():
    """Tambah, hapus, atau replace seluruh list RSS sources."""
    data = request.get_json(force=True)
    cfg = _load_scheduler_settings()
    sources = cfg.get("rss_sources", [])

    action = data.get("action", "replace")  # "add" | "remove" | "replace"

    if action == "replace":
        raw = data.get("rss_sources", [])
        sources = []
        for src in raw:
            u = (src.get("url") or "").strip()
            if u:
                sources.append({
                    "url": u,
                    "label": (src.get("label") or "").strip() or u,
                    "max_items": max(1, min(100, int(src.get("max_items") or 10))),
                })
    elif action == "add":
        u = (data.get("url") or "").strip()
        if not u:
            return jsonify({"error": "URL wajib diisi"}), 400
        # Cek duplikat URL
        if not any(s["url"] == u for s in sources):
            sources.append({
                "url": u,
                "label": (data.get("label") or "").strip() or u,
                "max_items": max(1, min(100, int(data.get("max_items") or 10))),
            })
    elif action == "remove":
        u = (data.get("url") or "").strip()
        sources = [s for s in sources if s["url"] != u]

    cfg["rss_sources"] = sources
    _save_scheduler_settings(cfg)
    return jsonify({"status": "ok", "rss_sources": sources})


@app.route("/api/scheduler/status", methods=["GET"])
def scheduler_status():
    cfg = _load_scheduler_settings()
    return jsonify({
        "enabled": cfg.get("enabled", False),
        "interval": cfg.get("interval", "manual"),
        "last_run_at": cfg.get("last_run_at"),
        "last_run_articles_added": cfg.get("last_run_articles_added", 0),
        "last_run_url": cfg.get("last_run_url", ""),
        "last_run_mode": cfg.get("last_run_mode", "full"),
        "next_run_at": _next_run_iso(),
        "scraper_running": scrape_state["running"],
    })


@app.route("/api/scheduler/run-now", methods=["POST"])
def scheduler_run_now():
    """Jalankan scheduled scrape sekarang juga (manual trigger)."""
    with state_lock:
        if scrape_state["running"]:
            return jsonify({"error": "Scraping sedang berjalan"}), 409
    cfg = _load_scheduler_settings()
    if not cfg.get("url", "").strip():
        return jsonify({"error": "URL scheduler belum dikonfigurasi"}), 400
    threading.Thread(target=_run_scheduled_scrape, daemon=True).start()
    return jsonify({"status": "started"})


def _log_startup_info():
    """Log informasi startup dan periksa env variables."""
    logger.info("=" * 60)
    logger.info("AINA News Scraper Backend — Starting Up")
    logger.info("=" * 60)

    if check_openai_available():
        logger.info("[ENV] OPENAI_API_KEY: tersedia ✓")
    else:
        logger.warning("[ENV] OPENAI_API_KEY: TIDAK DITEMUKAN — fitur AI Summary tidak aktif")

    if check_supabase_available():
        logger.info("[ENV] SUPABASE_URL + SUPABASE_SERVICE_ROLE_KEY: tersedia ✓")
    else:
        logger.warning("[ENV] SUPABASE_URL/SUPABASE_SERVICE_ROLE_KEY: TIDAK DITEMUKAN — fitur Push Supabase tidak aktif")

    logger.info(f"[ENV] Data dir: {os.path.abspath(DATA_DIR)}")
    logger.info("[STARTUP] Backend siap menerima request.")
    logger.info("=" * 60)


_log_startup_info()


# ─── Extra Sources ────────────────────────────────────────────────────────────

def _make_kb_draft(title: str, content: str, source_url: str = "",
                   published_date: str = "", source_tag: str = "") -> dict:
    """Helper: buat KB draft object dan simpan ke kb_articles.json."""
    import uuid as _uuid_m
    article_id = f"extra-{source_tag}-{int(time.time()*1000)}-{_uuid_m.uuid4().hex[:6]}"
    draft = convert_to_kb_format({
        "id": article_id,
        "title": title,
        "content": content,
        "url": source_url,
        "date": published_date or _now_iso(),
        "status": "success",
        "summary": "",
        "tags": [source_tag] if source_tag else [],
    })

    # ── Deteksi duplikat berdasarkan source_url ──
    kb = _load_kb()
    if source_url:
        for existing in kb:
            if existing.get("source_url", "").strip() == source_url.strip() and existing.get("id") != article_id:
                draft["is_duplicate"] = True
                draft["duplicate_of_id"] = existing.get("id", "")
                draft["duplicate_of_title"] = existing.get("title", "")
                draft["duplicate_of_status"] = existing.get("approval_status", "")
                break

    if not any(a.get("id") == article_id for a in kb):
        kb.append(draft)
        _save_kb(kb)
    return draft


@app.route("/api/youtube/scrape", methods=["POST"])
def api_youtube_scrape():
    """Ambil transkrip YouTube dan simpan sebagai KB draft."""
    data = request.get_json(force=True)
    url = (data.get("url") or "").strip()
    if not url:
        return jsonify({"error": "URL YouTube wajib diisi."}), 400

    # Ekstrak video ID dari berbagai format URL
    import re as _re
    vid_match = _re.search(
        r"(?:v=|youtu\.be/|/embed/|/v/|shorts/)([A-Za-z0-9_-]{11})", url
    )
    if not vid_match:
        return jsonify({"error": "URL YouTube tidak valid. Pastikan URL mengandung ID video (11 karakter)."}), 400
    video_id = vid_match.group(1)

    # Ambil judul video dari halaman YouTube
    title = f"YouTube Video {video_id}"
    try:
        import requests as _req
        r = _req.get(
            f"https://www.youtube.com/watch?v={video_id}",
            headers={"User-Agent": "Mozilla/5.0"},
            timeout=10,
        )
        og_match = _re.search(r'<meta property="og:title" content="([^"]+)"', r.text)
        if og_match:
            import html as _html
            title = _html.unescape(og_match.group(1))
    except Exception:
        pass

    # Ambil transkrip — pakai instance API (youtube-transcript-api >= 1.x)
    LANG_PREF = ["id", "en", "ar", "ms", "en-US", "en-GB"]
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        from youtube_transcript_api._errors import (
            TranscriptsDisabled, NoTranscriptFound, VideoUnavailable,
            IpBlocked, RequestBlocked, PoTokenRequired,
        )
        _yt_api = YouTubeTranscriptApi()
        raw = None
        try:
            transcript_list = _yt_api.list(video_id)
            # 1) Cari manual transcript dalam bahasa yang diinginkan
            try:
                transcript = transcript_list.find_transcript(LANG_PREF)
            except NoTranscriptFound:
                try:
                    # 2) Cari auto-generated transcript
                    transcript = transcript_list.find_generated_transcript(LANG_PREF)
                except NoTranscriptFound:
                    # 3) Ambil transcript pertama yang tersedia (bahasa apapun)
                    transcript = next(iter(transcript_list))
            raw = transcript.fetch()
        except (TranscriptsDisabled, VideoUnavailable, NoTranscriptFound):
            raise
        except (IpBlocked, RequestBlocked, PoTokenRequired):
            raise
        except Exception:
            # Fallback: coba langsung fetch multi-bahasa
            try:
                raw = _yt_api.fetch(video_id, languages=LANG_PREF)
            except Exception:
                raw = _yt_api.fetch(video_id)

        full_text = " ".join(s.text for s in raw)
        if not full_text.strip():
            return jsonify({"error": "Transkrip kosong atau tidak tersedia untuk video ini."}), 400

    except (IpBlocked, RequestBlocked, PoTokenRequired) as e:
        app.logger.warning(f"[YOUTUBE] IP/token block: {e}")
        return jsonify({"error": (
            "YouTube memblokir permintaan dari server ini. "
            "Coba lagi dalam beberapa menit, atau gunakan video lain."
        ), "hint": "ip_block"}), 503
    except TranscriptsDisabled:
        return jsonify({"error": (
            "Video ini tidak memiliki subtitle/CC yang aktif. "
            "Pemilik video menonaktifkan fitur subtitle, atau YouTube belum men-generate "
            "auto-subtitle untuk video ini."
        ), "hint": "no_cc"}), 400
    except VideoUnavailable:
        return jsonify({"error": "Video tidak tersedia — kemungkinan privat, dihapus, atau ID tidak valid."}), 400
    except NoTranscriptFound:
        return jsonify({"error": (
            "Tidak ditemukan transkrip untuk video ini dalam bahasa apapun "
            "(Indonesia, Inggris, Arab, Melayu)."
        ), "hint": "no_cc"}), 400
    except Exception as e:
        err = str(e)
        if "No transcripts" in err or "Could not retrieve" in err or "disabled" in err.lower():
            return jsonify({"error": "Video ini tidak memiliki subtitle/CC yang tersedia.", "hint": "no_cc"}), 400
        if "unavailable" in err.lower() or "not available" in err.lower():
            return jsonify({"error": "Video tidak tersedia atau bersifat privat."}), 400
        app.logger.error(f"[YOUTUBE] transcript error: {e}")
        return jsonify({"error": f"Gagal mengambil transkrip: {err}"}), 500

    draft = _make_kb_draft(title, full_text, url, source_tag="youtube")
    return jsonify({"status": "ok", "count": 1, "article": draft})


@app.route("/api/docx/upload", methods=["POST"])
def api_docx_upload():
    """Upload dan parse satu atau lebih file .docx/.doc menjadi KB drafts."""
    import tempfile as _tempfile, os as _os

    files = request.files.getlist("files")
    if not files:
        return jsonify({"error": "Tidak ada file yang diupload."}), 400

    try:
        import docx as _docx
        from docx.opc.exceptions import PackageNotFoundError as _PkgNotFound
    except ImportError:
        return jsonify({"error": "Library python-docx tidak tersedia."}), 500

    def _extract_text_from_doc(doc):
        """Ekstrak semua teks dari paragraf + sel tabel."""
        blocks = []
        # Paragraf biasa
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                blocks.append((para.style.name if para.style else "", t))
        # Teks di dalam tabel
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        t = para.text.strip()
                        if t and t not in [b[1] for b in blocks]:
                            blocks.append(("Normal", t))
        return blocks

    articles = []
    errors = []

    for f in files:
        fname = f.filename or "unknown.docx"
        fname_lower = fname.lower()
        tmp_path = None
        try:
            # Tolak .doc lama — python-docx hanya support .docx (Office Open XML)
            if fname_lower.endswith(".doc") and not fname_lower.endswith(".docx"):
                errors.append(
                    f"{fname}: Format .doc (Word lama) tidak didukung. "
                    "Buka di Word/LibreOffice lalu Save As → .docx, kemudian upload ulang."
                )
                continue

            # Simpan ke file temp
            tmp_fd, tmp_path = _tempfile.mkstemp(suffix=".docx")
            _os.close(tmp_fd)
            f.save(tmp_path)

            try:
                doc = _docx.Document(tmp_path)
            except _PkgNotFound:
                errors.append(
                    f"{fname}: File tidak bisa dibaca. Pastikan file adalah .docx yang valid "
                    "(bukan .doc lama, bukan file yang rusak/terenkripsi)."
                )
                continue

            blocks = _extract_text_from_doc(doc)

            # Cari judul: Title style, Heading 1, atau baris pertama yang pendek
            title = ""
            paragraphs = []
            for style_name, text in blocks:
                is_heading = style_name in ("Title", "Subtitle") or style_name.startswith("Heading")
                if not title and is_heading:
                    title = text
                else:
                    paragraphs.append(text)

            # Fallback: baris pertama pendek → jadikan judul
            if not title and paragraphs:
                if len(paragraphs[0]) <= 150:
                    title = paragraphs[0]
                    paragraphs = paragraphs[1:]

            # Fallback terakhir: nama file
            if not title:
                title = _os.path.splitext(fname)[0].replace("-", " ").replace("_", " ").title()

            content = "\n\n".join(paragraphs)
            if not content.strip():
                errors.append(f"{fname}: Tidak ada teks yang dapat diekstrak dari dokumen ini.")
                continue

            draft = _make_kb_draft(title, content, source_tag="docx")
            articles.append(draft)

        except Exception as e:
            errors.append(f"{fname}: {str(e)}")
        finally:
            if tmp_path and _os.path.exists(tmp_path):
                try:
                    _os.unlink(tmp_path)
                except Exception:
                    pass

    if not articles:
        err_detail = " | ".join(errors)
        return jsonify({"error": f"Tidak ada file yang berhasil diproses. {err_detail}"}), 400

    return jsonify({
        "status": "ok",
        "count": len(articles),
        "articles": articles,
        "errors": errors,
    })


@app.route("/api/rss/fetch", methods=["POST"])
def api_rss_fetch():
    """Fetch dan parse RSS/Atom feed, simpan setiap item sebagai KB draft."""
    data = request.get_json(force=True)
    url = (data.get("url") or "").strip()
    max_items = int(data.get("max_items") or 10)
    if not url:
        return jsonify({"error": "URL feed RSS wajib diisi."}), 400
    if max_items < 1:
        max_items = 10
    if max_items > 100:
        max_items = 100

    try:
        import feedparser as _fp
    except ImportError:
        return jsonify({"error": "Library feedparser tidak tersedia."}), 500

    try:
        feed = _fp.parse(url)
    except Exception as e:
        return jsonify({"error": f"Gagal fetch feed: {str(e)}"}), 500

    if feed.get("bozo") and not feed.get("entries"):
        return jsonify({"error": "URL bukan feed RSS/Atom yang valid atau tidak dapat diakses."}), 400

    entries = feed.get("entries", [])[:max_items]
    if not entries:
        return jsonify({"error": "Feed tidak memiliki entri / artikel."}), 400

    articles = []
    for entry in entries:
        title = entry.get("title", "").strip() or "(Tanpa Judul)"
        link = entry.get("link", "")
        # Coba content dulu, fallback ke summary
        content_raw = ""
        if entry.get("content"):
            content_raw = entry["content"][0].get("value", "")
        if not content_raw:
            content_raw = entry.get("summary", "") or entry.get("description", "")

        # Strip HTML dari content
        try:
            from bs4 import BeautifulSoup as _BS
            content = _BS(content_raw, "html.parser").get_text(separator="\n").strip()
        except Exception:
            import re as _re
            content = _re.sub(r"<[^>]+>", " ", content_raw).strip()

        if not content:
            content = title

        # Parse tanggal
        pub_date = ""
        if entry.get("published_parsed"):
            import time as _t
            try:
                from datetime import datetime as _dt
                pub_date = _dt(*entry["published_parsed"][:6]).strftime("%Y-%m-%dT%H:%M:%S")
            except Exception:
                pass

        draft = _make_kb_draft(title, content, link, pub_date, source_tag="rss")
        articles.append(draft)

    return jsonify({
        "status": "ok",
        "count": len(articles),
        "articles": articles,
        "feed_title": feed.feed.get("title", ""),
    })


@app.route("/api/telegram/scrape", methods=["POST"])
def api_telegram_scrape():
    """Scrape postingan dari Telegram channel publik via t.me/s/<channel>."""
    data = request.get_json(force=True)
    channel = (data.get("channel") or "").strip().lstrip("@")
    limit = int(data.get("limit") or 20)
    if not channel:
        return jsonify({"error": "Username channel Telegram wajib diisi."}), 400
    if limit < 1:
        limit = 20
    if limit > 200:
        limit = 200

    import requests as _req
    from bs4 import BeautifulSoup as _BS
    import re as _re

    preview_url = f"https://t.me/s/{channel}"
    try:
        r = _req.get(
            preview_url,
            headers={"User-Agent": "Mozilla/5.0 (compatible; AINA-Scraper/1.0)"},
            timeout=15,
        )
        if r.status_code == 404:
            return jsonify({"error": f"Channel @{channel} tidak ditemukan atau privat."}), 404
        if r.status_code != 200:
            return jsonify({"error": f"Gagal akses channel: HTTP {r.status_code}"}), 400
    except Exception as e:
        return jsonify({"error": f"Gagal terhubung ke Telegram: {str(e)}"}), 500

    soup = _BS(r.text, "html.parser")
    messages = soup.select(".tgme_widget_message_wrap")

    if not messages:
        return jsonify({"error": f"Tidak ada pesan publik yang ditemukan di @{channel}. Pastikan channel bisa diakses di t.me/s/{channel}"}), 400

    articles = []
    seen_texts: set = set()

    for msg in messages[:limit]:
        # Teks pesan
        text_el = msg.select_one(".tgme_widget_message_text")
        text = text_el.get_text(separator="\n").strip() if text_el else ""
        if not text or len(text) < 20:
            continue  # Skip pesan terlalu pendek (foto tanpa caption, dll)

        # Deduplicate
        text_key = text[:100]
        if text_key in seen_texts:
            continue
        seen_texts.add(text_key)

        # Tanggal
        time_el = msg.select_one(".tgme_widget_message_date time")
        pub_date = ""
        if time_el and time_el.get("datetime"):
            pub_date = time_el["datetime"][:19].replace(" ", "T")

        # Link ke pesan
        link_el = msg.select_one(".tgme_widget_message_date")
        msg_url = link_el.get("href", "") if link_el else ""

        # Judul: ambil baris pertama yang cukup panjang
        lines = [l.strip() for l in text.splitlines() if len(l.strip()) > 10]
        title = lines[0][:120] if lines else text[:80]

        draft = _make_kb_draft(title, text, msg_url or preview_url, pub_date, source_tag="telegram")
        articles.append(draft)

    if not articles:
        return jsonify({"error": "Tidak ada postingan teks yang berhasil diambil dari channel ini."}), 400

    return jsonify({
        "status": "ok",
        "count": len(articles),
        "articles": articles,
        "channel": channel,
    })


# ══════════════════════════════════════════════════════════════════════════════
# MUQARRAR — Upload PDF per-halaman, OCR, embed, simpan, tanya AINA
# ══════════════════════════════════════════════════════════════════════════════

_muqarrar_jobs: dict = {}   # job_id → {status, pages_done, pages_total, kitab_id, errors, kitab_name}

_LATIN_HEADINGS_MQ = re.compile(
    r'(?im)^(?:bab|pasal|fasal|bagian|chapter|section|unit|pelajaran|pertemuan|'
    r'tema|topik|materi|modul|pendahuluan|penutup|kesimpulan)\b[\s\d\w\-:\.]*$'
)
_ARAB_HEADINGS_MQ = re.compile(
    r'(?:الباب|الفصل|المبحث|المسألة|الفريضة|القسم|الموضوع|الدرس|الكتاب)'
)


def _embed_text(text: str, client) -> list[float] | None:
    """Single embed dengan retry. Wrapper ke _embed_texts_batch."""
    results = _embed_texts_batch([text], client)
    return results[0] if results else None


def _embed_texts_batch(texts: list, client, batch_size: int = 100, max_retries: int = 3) -> list:
    """
    Batch embed banyak teks sekaligus — JAUH lebih cepat dari satu-satu.
    200 halaman = 2 API call (batch 100), bukan 200 API call.
    Return list embedding (urutan sama dengan input). Gagal → [].
    """
    if not texts:
        return []

    all_embeddings = [[] for _ in texts]

    for start in range(0, len(texts), batch_size):
        batch = texts[start:start + batch_size]
        batch_inputs = [t[:8000] for t in batch]

        for attempt in range(max_retries):
            try:
                resp = client.embeddings.create(
                    model="text-embedding-3-small",
                    input=batch_inputs,
                    timeout=120,  # 2 menit per batch
                )
                # Susunan output bisa tidak berurutan — sort by index
                sorted_data = sorted(resp.data, key=lambda x: x.index)
                for i, item in enumerate(sorted_data):
                    all_embeddings[start + i] = item.embedding
                break  # sukses, lanjut batch berikutnya

            except Exception as e:
                wait = 2 ** attempt   # 1s, 2s, 4s
                if attempt < max_retries - 1:
                    logger.warning(
                        f"[MUQARRAR] Embedding batch [{start}:{start+len(batch)}] gagal "
                        f"(attempt {attempt+1}/{max_retries}), retry {wait}s: {e}"
                    )
                    time.sleep(wait)
                else:
                    logger.error(
                        f"[MUQARRAR] Embedding batch [{start}:{start+len(batch)}] "
                        f"gagal setelah {max_retries} percobaan: {e}"
                    )
                    # all_embeddings for this batch stay as []

    return all_embeddings


def _cosine_similarity(a: list, b: list) -> float:
    """Cosine similarity antara dua embedding vector."""
    va = np.array(a, dtype=np.float32)
    vb = np.array(b, dtype=np.float32)
    denom = np.linalg.norm(va) * np.linalg.norm(vb)
    if denom == 0:
        return 0.0
    return float(np.dot(va, vb) / denom)


def _page_to_image_bytes(fitz_doc, page_idx: int, dpi: int = 200) -> bytes:
    """Render halaman PDF ke PNG bytes untuk OCR."""
    page = fitz_doc[page_idx]
    pix = page.get_pixmap(dpi=dpi)
    return pix.tobytes("png")


def _detect_chapter_mq(text: str) -> str:
    """Deteksi heading bab/fasal dari teks halaman. Return heading string atau ''."""
    lines = text.strip().splitlines()[:6]
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if _LATIN_HEADINGS_MQ.match(line):
            return line[:120]
        if _ARAB_HEADINGS_MQ.search(line) and len(line) < 120:
            return line[:120]
    return ""


def _ocr_page_with_retry(fitz_doc, page_idx: int, page_num: int, client, max_retries: int = 3) -> str:
    """OCR satu halaman scan dengan retry dan timeout. Return teks atau ''."""
    import base64
    for attempt in range(max_retries):
        try:
            # Render halaman ke PNG — 150 DPI cukup untuk OCR teks (hemat RAM vs 200 DPI)
            img_bytes = _page_to_image_bytes(fitz_doc, page_idx, dpi=150)
            b64 = base64.b64encode(img_bytes).decode()
            del img_bytes  # bebaskan memori segera

            resp = client.chat.completions.create(
                model="gpt-4o",
                messages=[{
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                "Ini adalah halaman dari sebuah kitab/muqarrar. "
                                "Ekstrak SELURUH teks yang terlihat dengan akurat. "
                                "Untuk teks Arab: rekonstruksi huruf-huruf yang benar, pertahankan harakat jika ada. "
                                "Pertahankan struktur asli (judul, nomor, paragraf). "
                                "Output: HANYA teks yang diekstrak, tanpa komentar."
                            ),
                        },
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:image/png;base64,{b64}", "detail": "high"},
                        },
                    ],
                }],
                max_tokens=2500,
                temperature=0.0,
                timeout=120,   # 2 menit per halaman OCR
            )
            del b64  # bebaskan memori base64 setelah terpakai
            return resp.choices[0].message.content.strip()

        except Exception as e:
            wait = 3 ** attempt  # 1s, 3s, 9s
            if attempt < max_retries - 1:
                logger.warning(f"[MUQARRAR-OCR] hal {page_num} gagal (attempt {attempt+1}/{max_retries}), retry {wait}s: {e}")
                time.sleep(wait)
            else:
                logger.error(f"[MUQARRAR-OCR] hal {page_num} gagal setelah {max_retries} percobaan: {e}")
    return ""


def _process_muqarrar_upload(
    job_id: str,
    pdf_bytes: bytes,
    kitab_id: str,
    kitab_name: str,
    author: str,
    use_ocr_for_scans: bool,
    description: str = "",
):
    """
    Background worker — 3 Fase:
      Fase 1: Ekstrak teks per halaman (OCR jika scan)
      Fase 2: Batch embed semua teks sekaligus (200 hal = 2 API call)
      Fase 3: Simpan semua chunk ke Supabase
    """
    from ai_services import get_openai_client
    from db_services import muqarrar_save_chunk

    job = _muqarrar_jobs[job_id]
    try:
        client = get_openai_client()
        fitz_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        total = len(fitz_doc)
        job["pages_total"] = total
        job["status"] = "processing"
        job["phase"] = "extract"
        job["phase_label"] = "Ekstraksi teks"
        logger.info(f"[MUQARRAR] {job_id}: mulai proses {total} hal — {kitab_name}")

        # ══ FASE 1: Ekstrak teks per halaman ════════════════════════════════
        extracted: list = []   # [{page_num, text, chapter, word_count, is_ocr}]
        ocr_count = 0

        for page_idx in range(total):
            if job.get("cancelled"):
                logger.info(f"[MUQARRAR] {job_id}: dibatalkan oleh user")
                break

            page_num = page_idx + 1
            job["pages_done"] = page_idx
            job["current_page"] = page_num

            # Ekstrak teks digital
            fitz_page = fitz_doc[page_idx]
            raw_text = fitz_page.get_text("text").strip()
            is_ocr_page = False

            # Halaman scan / teks sangat tipis → coba OCR Vision
            if (not raw_text or len(raw_text) < 30) and use_ocr_for_scans:
                raw_text = _ocr_page_with_retry(fitz_doc, page_idx, page_num, client)
                is_ocr_page = bool(raw_text)
                if is_ocr_page:
                    ocr_count += 1
                job["ocr_count"] = ocr_count   # update live untuk frontend

            if not raw_text or len(raw_text) < 15:
                logger.debug(f"[MUQARRAR] hal {page_num}: skip (kosong)")
                continue

            chapter = _detect_chapter_mq(raw_text)
            word_count = len(raw_text.split())
            extracted.append({
                "page_num": page_num,
                "text": raw_text[:10000],
                "chapter": chapter,
                "word_count": word_count,
                "is_ocr": is_ocr_page,
            })

        fitz_doc.close()
        del pdf_bytes   # bebaskan memori PDF asli

        pages_extracted = len(extracted)
        logger.info(
            f"[MUQARRAR] Fase 1 selesai: {pages_extracted}/{total} hal diekstrak "
            f"({ocr_count} via OCR Vision)"
        )

        if not extracted:
            job["status"] = "error"
            job["error_msg"] = "Tidak ada teks yang berhasil diekstrak dari PDF."
            return

        # ══ FASE 2: Batch embed semua teks ══════════════════════════════════
        job["phase"] = "embed"
        job["phase_label"] = f"Membuat embedding ({pages_extracted} halaman)"
        job["pages_done"] = 0

        embed_inputs = []
        for item in extracted:
            prefix = f"Kitab: {kitab_name}\nHalaman: {item['page_num']}\n"
            if item["chapter"]:
                prefix += f"Bab/Fasal: {item['chapter']}\n"
            embed_inputs.append(prefix + "\n" + item["text"][:6000])

        # Batch embed — 200 hal → 2 API call (batch_size=100)
        embeddings = _embed_texts_batch(embed_inputs, client, batch_size=100, max_retries=3)
        logger.info(f"[MUQARRAR] Fase 2 selesai: {len(embeddings)} embeddings dibuat")

        # ══ FASE 3: Simpan semua chunk ke Supabase ═══════════════════════════
        job["phase"] = "save"
        job["phase_label"] = f"Menyimpan ke database ({pages_extracted} halaman)"
        job["pages_done"] = 0

        saved_count = 0
        for i, item in enumerate(extracted):
            if job.get("cancelled"):
                break

            job["pages_done"] = i
            chunk = {
                "id": f"{kitab_id}__p{item['page_num']}",
                "kitab_id": kitab_id,
                "kitab_name": kitab_name,
                "author": author,
                "description": description,
                "page_number": item["page_num"],
                "chapter": item["chapter"],
                "content": item["text"],
                "embedding": embeddings[i] if i < len(embeddings) else [],
                "word_count": item["word_count"],
                "is_ocr": item["is_ocr"],
            }

            # Retry simpan ke Supabase
            for attempt in range(3):
                if muqarrar_save_chunk(chunk):
                    saved_count += 1
                    break
                elif attempt < 2:
                    time.sleep(1)
                else:
                    job["errors"].append(f"Hal {item['page_num']}: gagal simpan")

        job["pages_done"] = pages_extracted
        job["pages_total"] = pages_extracted
        job["status"] = "done"
        job["kitab_id"] = kitab_id
        job["saved_count"] = saved_count
        job["ocr_count"] = ocr_count
        logger.info(
            f"[MUQARRAR] {job_id}: SELESAI — {kitab_name} | "
            f"{saved_count}/{pages_extracted} hal tersimpan | {ocr_count} OCR"
        )

    except Exception as e:
        logger.error(f"[MUQARRAR] {job_id}: error fatal — {e}", exc_info=True)
        job["status"] = "error"
        job["error_msg"] = str(e)


@app.route("/api/muqarrar/detect", methods=["POST"])
def api_muqarrar_detect():
    """
    Deteksi metadata kitab dari PDF menggunakan AI (GPT-4o).
    Baca halaman awal → ekstrak nama kitab, pengarang, deskripsi.
    """
    import json as _json
    from ai_services import get_openai_client, check_openai_available

    if not check_openai_available():
        return jsonify({"error": "OPENAI_API_KEY tidak ditemukan."}), 503

    if "file" not in request.files:
        return jsonify({"error": "Tidak ada file yang dikirim."}), 400

    pdf_file = request.files["file"]
    if not pdf_file.filename or not pdf_file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "File harus berformat PDF."}), 400

    try:
        pdf_bytes = pdf_file.read()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")

        # Ambil teks dari 8 halaman pertama
        pages_text = []
        for i in range(min(8, len(doc))):
            text = doc[i].get_text("text").strip()
            if text and len(text) > 20:
                pages_text.append(f"=== Halaman {i + 1} ===\n{text[:2500]}")
        doc.close()

        if not pages_text:
            return jsonify({"error": "Tidak ada teks yang bisa dibaca dari halaman awal PDF. Coba isi manual."}), 400

        combined = "\n\n".join(pages_text[:6])

        client = get_openai_client()
        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Kamu adalah asisten ekstraksi metadata kitab/buku Islam yang sangat teliti. "
                        "Analisis teks dari halaman awal sebuah PDF dan ekstrak informasi berikut. "
                        "Jawab HANYA dengan JSON valid, tanpa teks lain, tanpa markdown code block.\n"
                        'Format: {"kitab_name": "...", "author": "...", "description": "..."}\n\n'
                        "Panduan pengisian:\n"
                        "- kitab_name: nama lengkap kitab/buku paling tepat (sertakan nama Arab asli jika ada, "
                        "cth: 'Fath al-Qarib (فتح القريب)' atau 'Mabadi Fiqhiyyah Juz 1')\n"
                        "- author: nama lengkap pengarang/penulis beserta gelar jika ada "
                        "(kosong string jika tidak ditemukan)\n"
                        "- description: deskripsi informatif 2-3 kalimat dalam Bahasa Indonesia tentang "
                        "isi, tema, dan kegunaan kitab ini. Jadikan semenarik dan seinformatif mungkin."
                    ),
                },
                {
                    "role": "user",
                    "content": f"Berikut teks dari halaman awal PDF yang diupload:\n\n{combined}",
                },
            ],
            temperature=0.2,
            max_tokens=600,
        )

        raw = resp.choices[0].message.content.strip()
        # Hapus markdown code block jika ada
        if raw.startswith("```"):
            raw = raw.strip("`").strip()
            if raw.startswith("json"):
                raw = raw[4:].strip()

        result = _json.loads(raw)
        return jsonify({
            "kitab_name": (result.get("kitab_name") or "").strip(),
            "author": (result.get("author") or "").strip(),
            "description": (result.get("description") or "").strip(),
        })

    except _json.JSONDecodeError as e:
        logger.error(f"[MUQARRAR-DETECT] JSON parse error: {e}")
        return jsonify({"error": "AI memberikan respons tidak valid. Coba lagi."}), 500
    except Exception as e:
        logger.error(f"[MUQARRAR-DETECT] Error: {e}", exc_info=True)
        return jsonify({"error": f"Gagal mendeteksi: {str(e)}"}), 500


@app.route("/api/muqarrar/scan", methods=["POST"])
def api_muqarrar_scan():
    """
    Pra-scan PDF: baca struktur (jumlah halaman, TOC, deteksi bab)
    tanpa embedding — hanya PyMuPDF, cepat, tidak perlu AI.
    """
    import re as _re
    if "file" not in request.files:
        return jsonify({"error": "Tidak ada file yang dikirim."}), 400

    pdf_file = request.files["file"]
    if not pdf_file.filename or not pdf_file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "File harus berformat PDF."}), 400

    try:
        import fitz  # PyMuPDF
        pdf_bytes = pdf_file.read()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        pages_total = len(doc)

        # ── 1. Coba baca TOC native PDF ─────────────────────────────────────────
        raw_toc = doc.get_toc()  # [[level, title, page], ...]
        chapters = []
        toc_source = "none"

        if raw_toc and len(raw_toc) >= 1:
            toc_source = "native"
            for level, title, page in raw_toc:
                title = (title or "").strip()
                if not title:
                    continue
                chapters.append({
                    "level": level,
                    "title": title,
                    "page": max(1, page),
                })
            logger.info(f"[MUQARRAR SCAN] TOC native: {len(chapters)} entri")

        # ── 2. Jika TOC kosong / sangat sedikit → deteksi dari teks halaman ─────
        if len(chapters) < 3:
            toc_source = "detected"
            chapters = []

            # Pola deteksi: Latin + Arab + Melayu
            heading_patterns = [
                _re.compile(r'^(BAB|Bab|CHAPTER|Chapter|FASAL|Fasal|MUKADDIMAH|Mukaddimah|PENDAHULUAN|Pendahuluan|PENUTUP|Penutup|KESIMPULAN|Kesimpulan|DAFTAR|Daftar)\b', _re.IGNORECASE),
                _re.compile(r'^(فصل|باب|مقدمة|خاتمة|تمهيد|كتاب)', _re.UNICODE),
                _re.compile(r'^\d+[\.\-\)]\s+[A-ZА-Яا-ي]'),  # "1. Judul", "2- Bab"
                _re.compile(r'^(PART|Part|SECTION|Section|UNIT|Unit)\s+\w'),
            ]

            for page_idx in range(pages_total):
                page = doc[page_idx]
                page_num = page_idx + 1

                # Ambil teks dari blok teks terbesar (heading biasanya font besar)
                blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
                heading_found = None
                max_size = 0.0

                for blk in blocks:
                    if blk.get("type") != 0:  # type 0 = teks
                        continue
                    for line in blk.get("lines", []):
                        for span in line.get("spans", []):
                            text = (span.get("text") or "").strip()
                            size = span.get("size", 0)
                            if not text or len(text) > 120:
                                continue
                            # Cek pola heading
                            for pat in heading_patterns:
                                if pat.match(text):
                                    if size >= max_size:
                                        max_size = size
                                        heading_found = text
                                    break

                # Fallback: cek baris pertama teks halaman kalau font besar relatif
                if not heading_found:
                    plain = page.get_text("text").strip()
                    first_line = plain.split("\n")[0].strip() if plain else ""
                    if first_line and len(first_line) <= 80:
                        # Kalau font besar (ukuran relatif terhadap rata-rata)
                        all_sizes = [
                            sp.get("size", 0)
                            for blk in blocks if blk.get("type") == 0
                            for line in blk.get("lines", [])
                            for sp in line.get("spans", [])
                        ]
                        avg_size = (sum(all_sizes) / len(all_sizes)) if all_sizes else 10
                        if max_size > avg_size * 1.4 and any(pat.match(first_line) for pat in heading_patterns):
                            heading_found = first_line

                if heading_found:
                    chapters.append({
                        "level": 1,
                        "title": heading_found,
                        "page": page_num,
                    })

            logger.info(f"[MUQARRAR SCAN] Deteksi heading: {len(chapters)} bab terdeteksi")

        # ── 3. Hitung distribusi halaman per bab ────────────────────────────────
        for i, ch in enumerate(chapters):
            if i + 1 < len(chapters):
                ch["page_count"] = chapters[i + 1]["page"] - ch["page"]
            else:
                ch["page_count"] = pages_total - ch["page"] + 1

        # ── 4. Info halaman pertama (preview teks singkat) ──────────────────────
        first_page_preview = ""
        if pages_total > 0:
            fp_text = doc[0].get_text("text").strip()
            first_page_preview = fp_text[:300] if fp_text else ""

        doc.close()

        return jsonify({
            "pages_total": pages_total,
            "toc_source": toc_source,  # "native" | "detected" | "none"
            "chapters_count": len(chapters),
            "chapters": chapters,
            "first_page_preview": first_page_preview,
        })

    except Exception as e:
        logger.error(f"[MUQARRAR SCAN] error: {e}", exc_info=True)
        return jsonify({"error": f"Gagal scan PDF: {str(e)}"}), 500


@app.route("/api/muqarrar/db-status", methods=["GET"])
def api_muqarrar_db_status():
    """Periksa apakah tabel muqarrar_chunks sudah ada di Supabase."""
    from db_services import muqarrar_check_table
    result = muqarrar_check_table()
    return jsonify(result)


@app.route("/api/muqarrar/upload", methods=["POST"])
def api_muqarrar_upload():
    """Mulai proses upload & embedding muqarrar PDF. Return job_id."""
    from ai_services import check_openai_available
    if not check_openai_available():
        return jsonify({"error": "OPENAI_API_KEY tidak ditemukan."}), 503

    if "file" not in request.files:
        return jsonify({"error": "Tidak ada file yang dikirim."}), 400

    pdf_file = request.files["file"]
    if not pdf_file.filename or not pdf_file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "File harus berformat PDF."}), 400

    kitab_name = (request.form.get("kitab_name") or "").strip()
    author = (request.form.get("author") or "").strip()
    description = (request.form.get("description") or "").strip()
    use_ocr = request.form.get("use_ocr", "true").lower() == "true"

    if not kitab_name:
        return jsonify({"error": "Nama kitab wajib diisi."}), 400

    import time
    import uuid as _uuid
    import threading
    import re as _re

    # Buat kitab_id dari slug nama kitab + timestamp
    slug = _re.sub(r'[^a-z0-9]+', '-', kitab_name.lower()).strip('-')
    kitab_id = f"{slug}-{int(time.time())}"
    job_id = str(_uuid.uuid4())[:8]

    pdf_bytes = pdf_file.read()

    _muqarrar_jobs[job_id] = {
        "status": "queued",
        "phase": "queued",
        "phase_label": "Menunggu antrian",
        "pages_done": 0,
        "pages_total": 0,
        "current_page": 0,
        "kitab_id": kitab_id,
        "kitab_name": kitab_name,
        "saved_count": 0,
        "ocr_count": 0,
        "errors": [],
        "error_msg": "",
        "cancelled": False,
    }

    t = threading.Thread(
        target=_process_muqarrar_upload,
        args=(job_id, pdf_bytes, kitab_id, kitab_name, author, use_ocr, description),
        daemon=True,
    )
    t.start()

    return jsonify({"job_id": job_id, "kitab_id": kitab_id})


@app.route("/api/muqarrar/job/<job_id>", methods=["GET"])
def api_muqarrar_job(job_id: str):
    """Cek progress job upload muqarrar."""
    job = _muqarrar_jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job tidak ditemukan."}), 404
    return jsonify(job)


@app.route("/api/muqarrar/list", methods=["GET"])
def api_muqarrar_list():
    """Daftar semua kitab yang sudah diupload."""
    from db_services import muqarrar_list_kitab
    kitab_list = muqarrar_list_kitab()
    return jsonify({"kitab": kitab_list})


@app.route("/api/muqarrar/<kitab_id>", methods=["DELETE"])
def api_muqarrar_delete(kitab_id: str):
    """Hapus semua chunks satu kitab."""
    from db_services import muqarrar_delete_kitab
    ok = muqarrar_delete_kitab(kitab_id)
    if ok:
        return jsonify({"status": "ok", "deleted": kitab_id})
    return jsonify({"error": "Gagal menghapus kitab."}), 500


@app.route("/api/muqarrar/ask", methods=["POST"])
def api_muqarrar_ask():
    """
    Tanya AINA berdasarkan muqarrar yang telah diupload.
    Body: {question, kitab_id (opsional), top_k (opsional, default 5)}
    Return: {answer, sources: [{page, chapter, kitab_name, excerpt, score}]}
    """
    from ai_services import get_openai_client, check_openai_available
    from db_services import muqarrar_fetch_chunks_for_search

    if not check_openai_available():
        return jsonify({"error": "OPENAI_API_KEY tidak ditemukan."}), 503

    data = request.get_json(force=True) or {}
    question = (data.get("question") or "").strip()
    kitab_id = (data.get("kitab_id") or "").strip() or None
    top_k = min(int(data.get("top_k", 5)), 10)

    if not question:
        return jsonify({"error": "Pertanyaan tidak boleh kosong."}), 400

    try:
        client = get_openai_client()

        # ── 1. Embed pertanyaan ──────────────────────────────────────────────
        q_embedding = _embed_text(question, client)
        if not q_embedding:
            return jsonify({"error": "Gagal membuat embedding untuk pertanyaan."}), 500

        # ── 2. Fetch chunks dari Supabase ────────────────────────────────────
        chunks = muqarrar_fetch_chunks_for_search(kitab_id)
        if not chunks:
            return jsonify({"error": "Tidak ada data kitab yang ditemukan. Upload muqarrar terlebih dahulu."}), 404

        # ── 3. Hitung cosine similarity & ambil top-k ────────────────────────
        scored = []
        for c in chunks:
            emb = c.get("embedding")
            if not emb or not isinstance(emb, list) or len(emb) < 10:
                continue
            sim = _cosine_similarity(q_embedding, emb)
            scored.append((sim, c))

        scored.sort(key=lambda x: x[0], reverse=True)
        top_chunks = scored[:top_k]

        if not top_chunks:
            return jsonify({"error": "Tidak ada chunk yang relevan ditemukan."}), 404

        # ── 4. Susun konteks untuk GPT ───────────────────────────────────────
        context_parts = []
        sources_out = []
        for score, c in top_chunks:
            page = c["page_number"]
            chapter = c.get("chapter") or ""
            kname = c.get("kitab_name", "")
            author_str = c.get("author", "")
            content_preview = c["content"][:400].replace('\n', ' ').strip()

            label = f"[{kname}"
            if author_str:
                label += f" — {author_str}"
            label += f", Halaman {page}"
            if chapter:
                label += f", {chapter}"
            label += "]"

            context_parts.append(f"{label}\n{c['content'][:1500]}")
            sources_out.append({
                "page": page,
                "chapter": chapter,
                "kitab_name": kname,
                "author": author_str,
                "excerpt": content_preview,
                "score": round(score, 3),
            })

        context_text = "\n\n---\n\n".join(context_parts)

        # ── 5. Generate jawaban dengan sitasi halaman ────────────────────────
        system_msg = (
            "Kamu adalah AINA — asisten AI Islam berbasis pengetahuan kitab.\n\n"
            "INSTRUKSI KETAT:\n"
            "1. Jawab HANYA berdasarkan kutipan kitab yang diberikan di bawah.\n"
            "2. Setiap fakta/hukum yang kamu sebutkan WAJIB disertai sitasi: (Hal. X) atau (Hal. X-Y).\n"
            "3. Jika ada teks Arab yang relevan, kutip teksnya lalu berikan penjelasan.\n"
            "4. Jika informasi tidak cukup dari kutipan yang ada, katakan dengan jelas.\n"
            "5. Jangan tambahkan pengetahuan dari luar kutipan yang diberikan.\n"
            "6. Format jawaban: Markdown — gunakan **bold** untuk poin penting, "
            "`##` untuk sub-topik jika jawaban panjang.\n"
        )

        user_msg = (
            f"**Pertanyaan:** {question}\n\n"
            f"**Kutipan dari Kitab:**\n\n{context_text}\n\n"
            "Jawab pertanyaan di atas berdasarkan kutipan kitab. "
            "Sertakan nomor halaman untuk setiap informasi yang kamu ambil."
        )

        resp = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": user_msg},
            ],
            max_tokens=2000,
            temperature=0.1,
        )
        answer = resp.choices[0].message.content.strip()

        return jsonify({
            "status": "ok",
            "answer": answer,
            "sources": sources_out,
            "chunks_searched": len(chunks),
        })

    except Exception as e:
        logger.error(f"[MUQARRAR-ASK] Error: {e}")
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    # use_reloader=False: prevents Werkzeug stat reloader from restarting the process
    # mid-scrape (which would wipe in-memory scrape_state and cause progress log to blank out)
    app.run(host="0.0.0.0", port=8000, debug=True, use_reloader=False)
